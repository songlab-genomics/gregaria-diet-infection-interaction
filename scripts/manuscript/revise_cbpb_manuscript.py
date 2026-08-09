#!/usr/bin/env python3
"""Create a clean CBP Part B manuscript revision from the current Word draft.

The script never edits the source manuscript. It writes a dated revision, embeds
the current publication figures, and copies the proposed supplementary figures
to the same revision folder so the provenance of every panel remains explicit.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


PROJECT = Path(__file__).resolve().parents[2]
SOURCE = Path(
    "/Users/maevatecher/Desktop/"
    "Chapter # 3 Understanding the molecular mechanisms associated with diet and immune function (1).docx"
)
REVISION_DIR = PROJECT / "manuscript" / "revisions" / "20260808_cbpb_shared_infection_revision"
OUTPUT_DOCX = REVISION_DIR / "Chapter_3_CBPB_revised_20260808.docx"
OUTPUT_REVIEW_DOCX = REVISION_DIR / "Chapter_3_CBPB_changes_highlighted_20260808.docx"
SUPPLEMENT_DIR = REVISION_DIR / "supplementary_figures"

RUN_PUBLICATION = PROJECT / "output" / "rmd_runs" / "publication-figures_20260807-171511"
RUN_COUNT_SENSITIVITY = PROJECT / "output" / "rmd_runs" / "count-definition-sensitivity_20260808-000825"
RUN_FUNGAL = PROJECT / "output" / "rmd_runs" / "metarhizium-read-burden_20260807-143116"
RUN_CURATED = PROJECT / "output" / "rmd_runs" / "curated-target-gene-figure_20260808-001552"

FIGURES = {
    "figure1": RUN_PUBLICATION / "figure1_fatbody_pca_centroid_shifts.png",
    "figure2": RUN_PUBLICATION / "figure2_deg_burden_context_phase_venns.png",
    "figure3": RUN_PUBLICATION / "figure3_expression_cluster_gene_level_GO_KEGG_lollipop_eukaryote_inclusive.png",
    "figure4": RUN_CURATED / "figure4_curated_targets_main_chromosomes_descriptive.png",
    "figureS1": RUN_PUBLICATION / "figureS1_outlier_removal_sensitivity.png",
    "figureS2": RUN_PUBLICATION / "figureS2_deg_chromosome_localization.png",
    "figureS3": RUN_COUNT_SENSITIVITY / "count_definition_deg_totals.png",
    "figureS4": RUN_FUNGAL / "metarhizium_alignment_burden_combined.png",
    "figureS5": RUN_PUBLICATION / "figure4_curated_target_placement_impact.png",
    "figureS6": RUN_COUNT_SENSITIVITY / "count_definition_omnibus_interaction_LRT.png",
}


def clear_paragraph(paragraph):
    """Remove text and drawings while preserving the paragraph properties."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_text(paragraph, text: str, style: str | None = None, size: float | None = None):
    clear_paragraph(paragraph)
    if style is not None:
        paragraph.style = style
    run = paragraph.add_run(text)
    if size is not None:
        run.font.size = Pt(size)
    return paragraph


def blank(paragraph):
    clear_paragraph(paragraph)


def add_caption(paragraph, label: str, text: str):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(8)
    lead = paragraph.add_run(label + " ")
    lead.bold = True
    lead.font.size = Pt(9.5)
    body = paragraph.add_run(text)
    body.font.size = Pt(9.5)


def add_picture(paragraph, path: Path, width: float):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def add_supplement_before(target, image: Path, width: float, label: str, caption: str):
    image_paragraph = target.insert_paragraph_before()
    image_paragraph.paragraph_format.page_break_before = True
    add_picture(image_paragraph, image, width)
    caption_paragraph = target.insert_paragraph_before()
    add_caption(caption_paragraph, label, caption)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def validate_inputs():
    missing = [str(path) for path in [SOURCE, *FIGURES.values()] if not path.is_file()]
    if missing:
        raise FileNotFoundError("Missing manuscript input(s):\n" + "\n".join(missing))


def main():
    validate_inputs()
    REVISION_DIR.mkdir(parents=True, exist_ok=True)
    SUPPLEMENT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(SOURCE)
    p = doc.paragraphs

    # Correct the author display while preserving the existing superscript affiliations.
    for author_paragraph in (p[5], p[207]):
        for run in author_paragraph.runs:
            run.text = run.text.replace("Maeva A. Techera", "Maeva A. Techer")
            run.text = run.text.replace("Maeva Techera", "Maeva A. Techer")
            run.text = run.text.replace("Maeva Techer", "Maeva A. Techer")

    # The revised narrative targets CBP Part B rather than the omics-focused Part D.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(
                "Comparative Biochemistry and Physiology Part D: Genomics and Proteomics",
                "Comparative Biochemistry and Physiology Part B: Biochemistry and Molecular Biology",
            )
            run.text = run.text.replace(
                "Comparative Biochemistry and Physiology, Part D: Genomics and Proteomics",
                "Comparative Biochemistry and Physiology, Part B: Biochemistry and Molecular Biology",
            )

    # Submission highlights: each statement remains under 100 characters.
    replace_text(p[27], "Fungal infection strongly reshaped the desert locust fat body transcriptome")
    replace_text(p[28], "Infection-responsive expression was broadly shared across nutritional states")
    replace_text(p[29], "A core of 467 infection-responsive genes was shared among all three diets")
    replace_text(p[30], "Formal interaction tests supported selective, count-sensitive diet effects")

    abstract = (
        "Dietary macronutrient balance alters disease outcome in desert locusts, yet carbohydrate-biased "
        "nutrition can improve survival despite weaker phenoloxidase responses. We tested two competing "
        "explanations for this physiological paradox: enhanced immune investment under carbohydrate-biased "
        "nutrition versus diet-dependent modification of the host metabolic environment during infection. "
        "Fat body transcriptomes were measured from 44 fifth-instar female Schistocerca gregaria assigned to "
        "carbohydrate-biased (diet 33), balanced (diet 50), or protein-biased (diet 83) diets and challenged "
        "with Metarhizium robertsii or an oil control. Infection altered 1,016, 2,009, and 1,030 genes within "
        "diets 33, 50, and 83, respectively, with 467 genes shared among all three responses. Immune recognition, "
        "extracellular defense, transport, and Toll/Imd-associated functions were induced across nutritional "
        "conditions, whereas cuticle, cell-cycle, and selected metabolic functions generally decreased. A formal "
        "diet-by-infection likelihood-ratio test identified only nine genes in the primary transcript-plus-exon "
        "analysis and none under exon-only counting, indicating limited and count-definition-sensitive interaction. "
        "Post hoc curated immune candidates were commonly induced across diets, whereas protein-anabolism, nutrient "
        "sensing, amino-acid metabolism, and storage candidates showed heterogeneous responses. These results do not "
        "support generalized enhancement of immune transcription as the explanation for carbohydrate-associated "
        "survival. Instead, they are consistent with selective effects of nutritional state on host metabolism, "
        "resource availability, or tolerance, while leaving the causal mechanism unresolved."
    )
    replace_text(p[45], abstract)

    replace_text(p[55], "2. Materials and methods", "Heading 1")
    replace_text(p[57], "2.1 Experimental design", "Heading 3")
    replace_text(
        p[58],
        "We used a 3 x 2 factorial design comprising three isocaloric diets and two inoculation treatments. "
        "The diets contained 33%, 50%, or 83% protein as a proportion of total dietary macronutrient, "
        "corresponding to whole-diet protein:carbohydrate percentages of 14:28, 21:21, and 35:7, respectively. "
        "Individual fifth-instar female Schistocerca gregaria were randomly assigned to a diet and received "
        "either Metarhizium robertsii conidia in oil or sterile oil alone. The RNA-seq data set analyzed here "
        "comprised 44 individuals: diet 33 included seven controls and nine infected locusts, diet 50 included "
        "six controls and nine infected locusts, and diet 83 included seven controls and six infected locusts "
        "(20 controls and 24 infected locusts in total). All 44 libraries were retained in the primary analysis."
    )

    replace_text(p[60], "2.2 Host, pathogen, and diets", "Heading 3")
    replace_text(p[61], "Study animals", "Heading 4")
    replace_text(
        p[62],
        "Fifth (final) instar female S. gregaria, three days post-molt, were obtained from a gregarious colony "
        "maintained at Arizona State University from laboratory lines originating at the University of "
        "Leicester (United Kingdom) and the University of Haifa-Oranim (Israel). Age was controlled during "
        "density rearing to reduce age-dependent transcriptional variation, and the three-day interval allowed "
        "cuticular hardening before inoculation. Colony conditions were 34 degrees C during the day, 25 degrees C "
        "at night, 20-50% relative humidity, and a 14 h:10 h light:dark cycle. Colony locusts received organic "
        "romaine lettuce, wheat grass, and wheat bran. Experimental locusts were housed individually in "
        "ventilated plastic containers (18.9 x 13.5 x 9.5 cm) with food and water provided ad libitum."
    )
    replace_text(p[63], "Fungal culture and inoculation", "Heading 4")
    replace_text(
        p[64],
        "Locusts were challenged with M. robertsii isolate DWR2009 (also designated ARSEF 10343), originally "
        "isolated from soil in the United States. Conidia were produced by biphasic solid-state fermentation "
        "following Jaronski and Jackson (2012), suspended in Orchex 792 paraffinic oil, and adjusted to 1 x "
        "10^9 conidia mL-1. Concentration was determined with a hemocytometer and viability was confirmed by "
        "germination on potato dextrose agar at 25-26 degrees C. Infected locusts received 1 microliter of the "
        "conidial suspension at the hind-leg femur-trochanter junction using a Hamilton Gastight syringe with "
        "a repeating dispenser and blunt-ended PTFE needle. Controls received 1 microliter of sterile Orchex "
        "792 oil and were handled identically."
    )
    replace_text(p[65], "Artificial diets", "Heading 4")
    replace_text(
        p[66],
        "Chemically defined diets followed Dadd (1961), Simpson and Abisgold (1985), and the formulation used "
        "by Zembrzuski et al. (2023). Each diet contained 42% macronutrient by dry mass and was supplemented "
        "with Wesson's salts (2.4%), cholesterol (0.5%), linoleic acid (0.5%), ascorbic acid (0.3%), and a "
        "multivitamin mixture (0.2%). Protein was supplied as a 3:1:1 mixture of casein, peptone, and albumen; "
        "carbohydrate was supplied as equal parts sucrose and dextrin. Diet and water were replenished as needed."
    )

    replace_text(p[68], "2.3 Tissue sampling and RNA extraction", "Heading 3")
    blank(p[69])
    replace_text(
        p[70],
        "Fat body was collected 96 h after inoculation because this tissue integrates insect immunity, "
        "nutrient storage, and intermediary metabolism. A single experimenter performed all dissections under "
        "sterile, RNase-free conditions. Instruments were flame sterilized and sequentially treated with 10% "
        "bleach, distilled water, 70% ethanol, and RNaseZap before each dissection. Locusts were surface "
        "sterilized with 70% ethanol, and abdominal fat body was removed while avoiding the gut, trachea, and "
        "heart. Tissue was transferred immediately to nuclease-free tubes, flash-frozen in liquid nitrogen, "
        "and stored at -80 degrees C."
    )
    replace_text(
        p[71],
        "Total RNA was extracted on a Promega Maxwell RSC instrument with the simplyRNA Tissue Kit according "
        "to the manufacturer's protocol. RNA concentration was measured with a Qubit 4 Fluorometer and the "
        "Qubit RNA HS Assay. RNA integrity and dissection quality were assessed before library preparation."
    )

    replace_text(p[72], "2.4 Library preparation and sequencing", "Heading 3")
    replace_text(
        p[73],
        "Stranded total-RNA libraries were prepared at Texas A&M University with the Illumina Stranded Total "
        "RNA kit and Ribo-Zero rRNA depletion. Pooled libraries were sequenced by Novogene on an Illumina "
        "NovaSeq X Plus 25B platform as 150-bp paired-end reads, targeting approximately 45 million read pairs "
        "per library."
    )
    replace_text(
        p[74],
        "Residual rRNA was quantified from assigned gene-level fragments before filtering. The median residual "
        "rRNA fraction was 7.9% per library (range 3.3-41.7%). All 12,402 gene identifiers in the project rRNA "
        "exclusion list were removed before expression filtering, normalization, ordination, and differential "
        "expression testing."
    )

    replace_text(p[75], "2.5 Read processing, competitive mapping, and gene counting", "Heading 3")
    replace_text(
        p[76],
        "RNA-seq preprocessing was executed on a high-performance computing cluster with a run-scoped "
        "Snakemake workflow. Paired reads were processed with fastp using automatic paired-end adapter "
        "detection, removal of the first two bases of each read, and a minimum retained length of 50 bp. JSON "
        "and HTML quality reports were retained for each library."
    )
    replace_text(
        p[77],
        "Trimmed reads were aligned in two-pass mode with STAR v2.7.10b. The primary branch used a competitive "
        "reference comprising the S. gregaria assembly GCF_023897955.1 and the M. robertsii ARSEF 23 assembly "
        "GCF_000187425.2. The latter is a close reference genome rather than the exact challenge isolate "
        "DWR2009/ARSEF 10343. A host-only alignment branch was retained as a mapping sensitivity analysis. "
        "Genome indices used sjdbOverhang = 149 and alignIntronMax = 2,500,000, selected after examination of "
        "Schistocerca intron lengths. Coordinate-sorted BAMs and transcriptome BAMs were retained and indexed "
        "with samtools v1.21."
    )
    replace_text(
        p[78],
        "Gene-level fragments were quantified with featureCounts (Subread) using paired-end fragment counting "
        "(--countReadPairs), feature types transcript and exon (-t transcript,exon), and gene_id as the grouping "
        "attribute (-g gene_id). This span-inclusive definition assigns fragments overlapping either annotated "
        "feature class to a gene and can capture signal consistent with incompletely processed or nascent RNA; "
        "an exon-only matrix was analyzed separately as a count-definition sensitivity check."
    )
    replace_text(
        p[79],
        "Sequence placement was assigned from the official assembly metadata rather than scaffold-name patterns. "
        "The primary manuscript analysis was restricted to non-rRNA genes on the main assembled chromosomes. "
        "All-scaffold results were retained as a placement sensitivity analysis, and competitive-host and "
        "host-only placed-chromosome DEG totals were compared to assess mapping dependence."
    )

    replace_text(p[80], "2.6 Differential expression and sensitivity analyses", "Heading 3")
    replace_text(
        p[81],
        "DESeq2 was applied to raw gene counts after rRNA and sequence-placement filtering. Genes entered a model "
        "when they had at least 10 counts in at least three libraries. Automatic count replacement was disabled "
        "(minReplicatesForReplace = Inf), and Cook's-distance filtering was disabled for planned contrasts "
        "(cooksCutoff = FALSE), so all 44 samples contributed to the primary estimates. A gene was called "
        "differentially expressed at a Benjamini-Hochberg adjusted P value below 0.05 and an absolute log2 fold "
        "change of at least 1. Variance-stabilized counts were used for PCA and heatmaps, not for DESeq2 testing."
    )
    replace_text(
        p[82],
        "A six-level diet-by-treatment group model (~ 0 + group) generated direct within-diet infection contrasts "
        "and diet contrasts separately among controls and infected locusts. A second model (~ Diet + Treatment) "
        "estimated the average infection effect across diets. Diet-dependent infection responses were evaluated "
        "first with an omnibus likelihood-ratio test comparing a full model (~ Diet + Treatment + Diet:Treatment) "
        "with a reduced additive model (~ Diet + Treatment), and then with three pairwise difference-of-differences "
        "contrasts. Omnibus genes were called at a Benjamini-Hochberg FDR below 0.05; because the LRT tests multiple "
        "interaction coefficients jointly, effect size was summarized using the largest absolute pairwise "
        "interaction log2 fold change. For diet contrasts, diet 50 "
        "served as the reference for comparisons with diets 33 and 83, and diet 33 served as the reference for "
        "the 83-versus-33 comparison. Positive log2 fold changes indicate higher expression in the named "
        "numerator condition."
    )
    replace_text(
        p[83],
        "Library size, variance-stabilized PCA, sample-distance heatmaps, and heatmaps of highly variable genes "
        "were examined for sample-level quality control. Centroids were calculated for each diet-by-treatment "
        "group, and lines joining control and infected centroids summarized the direction of the infection shift "
        "within each diet."
    )
    replace_text(
        p[84],
        "The primary analysis retained every library. A supplementary sample-retention sensitivity analysis "
        "repeated the principal DEG contrasts after removing samples 1007, 1036, 1037, and 1039, which had been "
        "identified during exploratory ordination as having atypical global expression profiles. Because these "
        "libraries passed sequencing and mapping quality controls and no technical cause was established, this "
        "secondary branch was used only to evaluate robustness. Additional sensitivities compared exon-only with "
        "transcript-plus-exon counts and main-chromosome with all-scaffold analyses. None of these sensitivity "
        "branches replaced the all-sample, transcript-plus-exon, main-chromosome primary analysis."
    )

    replace_text(p[86], "2.7 Gene classification and functional enrichment", "Heading 3")
    replace_text(
        p[87],
        "Reference annotation and local crosswalk tables classified genes as protein coding, long non-coding RNA "
        "(lncRNA), or other/unknown. Because every listed rRNA locus was removed before modeling, the latter "
        "category contains non-rRNA features or genes without a resolved biotype."
    )
    replace_text(
        p[88],
        "Functional descriptions, Gene Ontology (GO) terms, and KEGG orthology/pathway assignments were obtained "
        "from the project eggNOG annotation and linked to gene identifiers through CDS records in the reference "
        "GFF. KEGG results were shown in a eukaryote-inclusive view and were interpreted as orthology-based, "
        "hypothesis-generating annotations rather than direct evidence that an insect pathway was activated."
    )
    replace_text(
        p[89],
        "Over-representation tests were performed with clusterProfiler::enricher using the 14,132 expressed, "
        "non-rRNA main-chromosome genes entering the primary DESeq2 analysis as the background. P values were "
        "adjusted by the Benjamini-Hochberg method. For the two expression clusters, enrichment foregrounds were "
        "restricted to genes differentially expressed in at least one within-diet infected-versus-control "
        "contrast. GO Biological Process, Cellular Component, and Molecular Function terms and eukaryote-compatible "
        "KEGG pathways were summarized separately. Term-to-gene reverse-index tables were retained so each "
        "enrichment result could be traced to its contributing genes."
    )

    replace_text(p[91], "2.8 Curated physiological candidate genes", "Heading 3")
    replace_text(
        p[92],
        "A post hoc curated candidate list was used to describe immune recognition and signaling, humoral immunity, cellular "
        "immunity, detoxification, protein synthesis and translation, nutrient sensing and TOR/insulin signaling, "
        "nitrogen and amino-acid metabolism, and storage hexamerins. The primary placed-chromosome analysis "
        "retained 99 candidates. Genes had originally entered this list when they were significant in at least one "
        "examined comparison; consequently, the displayed proportions were treated as descriptive summaries rather "
        "than independent category-level hypothesis tests. For each contrast, candidates were classified as "
        "upregulated, downregulated, or not significant using the same DEG thresholds as the genome-wide analysis. "
        "The all-scaffold candidate set was retained as a placement sensitivity analysis."
    )

    replace_text(
        p[90],
        "Genes significant in any primary contrast were clustered from row-standardized variance-stabilized "
        "expression using Euclidean distance and complete-linkage hierarchical clustering. Candidate cuts from "
        "k = 2 to k = 10 were compared by silhouette width; k = 2 had the highest mean silhouette width and defined "
        "the two primary expression programs. Genes were ordered within each cluster without averaging individual-"
        "gene expression, preserving sample-level heterogeneity in the heatmap. Cluster enrichment used within-diet "
        "infection DEGs as foregrounds and the expressed, rRNA-filtered main-chromosome genes as the universe."
    )

    # Add the phase-reference specificity check and fungal-RNA proxy before Results.
    phase_heading = p[94].insert_paragraph_before(
        "2.9 Phase-reference specificity analysis", style="Heading 3"
    )
    phase_text = p[94].insert_paragraph_before(
        "Because experimental locusts were housed individually for 96 h, the union of fat-body genes significant "
        "in at least one within-diet infection contrast was compared with independent solitarious-versus-gregarious "
        "DEG sets from S. gregaria head and thorax. Gene identifiers were matched directly and set intersections "
        "were summarized without re-estimating expression across tissues. These reference data came from different "
        "tissues and experimental contexts; the comparison was therefore used only to assess whether the dominant "
        "infection response broadly recapitulated known density-associated expression, not to infer a phase change "
        "in the experimental fat body."
    )
    phase_heading.paragraph_format.keep_with_next = True
    phase_text.paragraph_format.space_after = Pt(8)

    fungal_heading = p[94].insert_paragraph_before("2.10 Fungal alignment proxy", style="Heading 3")
    fungal_text = p[94].insert_paragraph_before(
        "Uniquely aligned fragments assigned to the fungal component of the competitive reference were summarized "
        "as counts and as a percentage of uniquely aligned competitive-reference fragments. Because the mapping "
        "reference was not strain matched and fungal load was not confirmed post mortem by qPCR or culture, this "
        "quantity was treated only as RNA-seq evidence of fungal transcriptional signal. A parallel qPCR validation "
        "would have been preferable, but no matching tissue remained for that assay."
    )
    fungal_heading.paragraph_format.keep_with_next = True
    fungal_text.paragraph_format.space_after = Pt(8)

    # Results are rewritten around the current placed-chromosome outputs.
    replace_text(p[95], "3.1 Fungal infection dominates fat body expression across diets", "Heading 3")
    replace_text(
        p[96],
        "The first two principal components of variance-stabilized expression explained 27.7% and 12.4% of total "
        "variance, respectively (Fig. 1). Control samples occupied the negative portion of PC1, whereas most "
        "infected samples shifted toward positive PC1 values. The control-to-infected centroid moved in the same "
        "general direction for each diet, although the hulls overlapped and within-group dispersion differed, "
        "particularly among infected locusts. Infection therefore represented the dominant common axis of "
        "variation, with diet and inter-individual heterogeneity modifying the position and breadth of each group."
    )
    replace_text(
        p[97],
        "Fungal alignment signal was low in most fat-body libraries but highly heterogeneous among infected "
        "individuals. Median fungal percentages were 0.00138%, 0.00097%, and 0.00084% in infected diets 33, 50, "
        "and 83, respectively, while a small number of infected libraries reached 0.64-3.13%. Controls remained "
        "near background. Because these values came from competitive RNA-seq alignment rather than a strain-matched "
        "load assay, they were used to describe detectable fungal RNA, not to classify infection success."
    )
    add_picture(p[98], FIGURES["figure1"], 6.2)
    add_caption(
        p[99],
        "Figure 1.",
        "Fat body transcriptomes shift with fungal infection across the three diets. PCA used variance-stabilized "
        "counts from expressed, non-rRNA genes on the main S. gregaria chromosomes. Points represent individual "
        "libraries, shapes indicate treatment, colors indicate diet, shaded polygons show the observed breadth of "
        "each diet-by-treatment group, open circles mark centroids, and lines connect the control and infected "
        "centroids within each diet."
    )

    replace_text(p[100], "3.2 Differential expression across infection and diet contrasts", "Heading 3")
    replace_text(
        p[101],
        "The average infected-versus-control contrast identified 1,691 DEGs across diets, including 893 genes "
        "higher in infected locusts and 798 higher in controls (Fig. 2). Within diets, infection altered 1,016 "
        "genes in diet 33 (552 higher in infected and 464 higher in controls), 2,009 genes in diet 50 (1,263 and "
        "746), and 1,030 genes in diet 83 (658 and 372). Protein-coding genes formed the majority of each response, "
        "but lncRNAs contributed 51, 132, and 62 DEGs in diets 33, 50, and 83, respectively."
    )
    replace_text(
        p[102],
        "Direct diet effects were much smaller among controls: pairwise contrasts yielded 15 DEGs between diets "
        "33 and 50, 60 between diets 83 and 33, and 59 between diets 83 and 50. Diet contrasts among infected "
        "locusts yielded 303, 86, and 59 DEGs for the same comparisons. The particularly large infected 33-versus-50 "
        "contrast was strongly directional, with 290 genes higher in diet 50 and 13 higher in diet 33. Across the "
        "three within-diet infection responses, 467 genes were shared by all diets, whereas 211, 1,069, and 320 "
        "were unique to diets 33, 50, and 83, respectively."
    )
    replace_text(
        p[103],
        "Formal interaction testing indicated that most of the infection response was shared rather than broadly "
        "diet specific. The omnibus diet-by-infection LRT identified nine genes in the primary transcript-plus-exon "
        "main-chromosome analysis; all nine had a maximum pairwise interaction effect of at least |log2FC| = 1. "
        "Pairwise difference-of-differences contrasts identified 0, 58, and 11 DEGs for diets 33 versus 50, 83 "
        "versus 50, and 83 versus 33, respectively, and all nine omnibus genes appeared in at least one of the two "
        "contrasts involving diet 83. Under exon-only counting, no gene passed the omnibus interaction FDR. The "
        "evidence for diet-dependent modification was therefore selective and sensitive to count definition, "
        "despite the marked differences in within-diet DEG totals."
    )
    add_picture(p[104], FIGURES["figure2"], 6.15)
    add_caption(
        p[105],
        "Figure 2.",
        "Differential-expression burden and overlap in the all-sample, main-chromosome analysis. The upper panels "
        "show coding, lncRNA, and other/unknown DEGs by direction for the average infection response, infection "
        "within each diet, diet contrasts among controls, and diet contrasts among infected locusts. Red values are "
        "higher in the named numerator condition and blue values are higher in the denominator. The lower-left Venn "
        "diagram partitions the three within-diet infection DEG sets. The lower-right diagram compares their union "
        "with external head and thorax solitarious-versus-gregarious DEG sets; this overlap provides phase-related "
        "specificity context but does not establish phase effects in fat body. Of 2,594 infection-responsive genes "
        "in this union, 2,134 did not overlap either reference phase set. Annotated rRNA genes were absent from "
        "every set."
    )
    blank(p[106])
    blank(p[107])

    replace_text(p[110], "3.3 Infection-associated expression programs and functional enrichment", "Heading 3")
    replace_text(
        p[111],
        "The union of genes significant in the primary contrast set comprised 2,949 main-chromosome genes. "
        "Complete-linkage clustering of Euclidean expression distances resolved two broad programs without "
        "averaging gene-level expression (Fig. 3); the two-cluster cut had the strongest mean silhouette width "
        "among candidate cuts from k = 2 to k = 10. "
        "Cluster 1 contained 1,086 genes and generally showed higher expression in controls and lower expression "
        "after infection. Cluster 2 contained 1,863 genes and showed the reciprocal pattern. Most genes in both "
        "clusters were associated with an infection-within-diet contrast; smaller subsets reflected diet effects "
        "or membership in more than one contrast context."
    )
    replace_text(
        p[112],
        "Among within-diet infection DEGs, cluster 1 was enriched for chitin-based cuticle development, mitotic "
        "cell-cycle and spindle functions, oxidoreductase activity, broad metabolism, cytochrome P450-associated "
        "xenobiotic metabolism, and insect hormone biosynthesis. Cluster 2 was enriched for response to bacteria "
        "and stimulus, extracellular and plasma-membrane compartments, peptidoglycan metabolism and binding, "
        "transmembrane transport, and Toll and Imd signaling. These reciprocal programs are consistent with "
        "reduced investment in cuticle, cell-cycle, and selected metabolic processes alongside increased expression "
        "of extracellular defense and recognition functions during infection."
    )
    replace_text(
        p[113],
        "Only 325 cluster-1 and 396 cluster-2 foreground genes had GO annotations, and 264 and 318, respectively, "
        "had KEGG assignments. The enrichment results therefore summarize the annotated subset rather than every "
        "DEG. Fungal alignment percentages displayed above the heatmap showed that the host expression "
        "programs occurred across a wide range of fungal RNA signals, but this visual correspondence was not "
        "interpreted as a validated relationship with pathogen load."
    )
    add_picture(p[108], FIGURES["figure3"], 5.15)
    add_caption(
        p[109],
        "Figure 3.",
        "Main-chromosome DEG expression programs and functional enrichment. (A) Gene-level row-z-scored "
        "variance-stabilized expression for 2,949 DEGs, separated by complete-linkage Euclidean clustering into "
        "the two-cluster cut with the strongest mean silhouette width and "
        "ordered with controls before infected samples. The upper bar shows the percentage of unique competitive "
        "alignments assigned to M. robertsii; values above the 0.1% display range are labeled. (B) GO terms enriched "
        "among within-diet infection DEGs in each cluster. (C) Eukaryote-inclusive KEGG pathways enriched in the "
        "same foregrounds. Point size indicates the number of cluster genes in a term. Cluster-1 enrichment values "
        "are mirrored to the left only to separate the panels; the sign does not represent regulation direction."
    )

    replace_text(p[114], "3.4 Curated immune and physiological candidate genes", "Heading 3")
    replace_text(
        p[115],
        "The main-chromosome curated analysis retained 99 genes across eight immune and protein-anabolism "
        "categories (Fig. 4). Within each diet, infection increased a substantial fraction of genes involved in "
        "innate immune recognition and signaling and humoral immunity. Responses in cellular immunity and "
        "detoxification were more variable, and several genes in these small categories were not significant."
    )
    replace_text(
        p[116],
        "Protein-anabolism categories contained both increased and decreased responses. Protein synthesis and "
        "translation included a small upregulated fraction and a larger downregulated fraction in every diet. "
        "Nutrient-sensing and nitrogen/amino-acid categories tended to contain more upregulated genes in diet 50 "
        "than in diets 33 or 83, whereas storage hexamerins were mostly downregulated in diet 50 and not significant "
        "in diets 33 and 83."
    )
    replace_text(
        p[117],
        "Across-diet contrasts among infected locusts contained relatively few significant curated genes. Because "
        "the curated list had been assembled from genes significant in at least one comparison, its status "
        "proportions were not subjected to an independent category-level significance test. The figure is therefore "
        "a descriptive map of infection-responsive physiological candidates and identifies categories for future "
        "testing with an annotation-defined, a priori gene set."
    )
    replace_text(
        p[118],
        "Restricting the curated list to main chromosomes reduced it from 163 all-scaffold candidates to 99. This "
        "placement filter was applied consistently to the primary PCA, differential-expression, clustering, "
        "enrichment, and curated analyses; the all-scaffold results are retained as a supplementary sensitivity."
    )
    blank(p[119])
    add_picture(p[120], FIGURES["figure4"], 6.15)
    blank(p[121])
    add_caption(
        p[122],
        "Figure 4.",
        "Curated immune and protein-anabolism candidate genes on the main S. gregaria chromosomes. Stacked bars "
        "show the percentage of genes that were upregulated, downregulated, or not significant. (A, C) Infection "
        "responses within each diet. (B, D) Pairwise diet contrasts among infected locusts, oriented as higher- "
        "versus lower-protein diets. The number below each functional heading is the number of retained candidate "
        "genes. Candidate membership was selected post hoc from genes significant in at least one contrast, so this "
        "panel is descriptive and does not present independent category-level p-values."
    )
    blank(p[123])

    # Discussion: retain the manuscript's voice while removing claims not supported by the current tests.
    replace_text(
        p[125],
        "Metarhizium infection was the principal driver of transcriptional variation in the desert locust fat "
        "body. Control-to-infected centroids moved in a common direction across diets, the average infection "
        "contrast identified 1,691 DEGs, and 467 genes were shared by all three within-diet infection responses. "
        "The balanced diet produced the largest descriptive DEG total, but the omnibus interaction identified only "
        "nine genes under the primary transcript-plus-exon definition and none under exon-only counting. Different "
        "numbers of significant genes among diets therefore should not be interpreted as evidence of broad "
        "diet-dependent interaction. Instead, the data support an extensive shared infection programme with "
        "selective and count-definition-sensitive dietary modification."
    )
    replace_text(
        p[126],
        "The two expression programs provide a physiological summary of this common response. One programme "
        "decreased during infection "
        "and was associated with cuticle, cell-cycle, and selected metabolic processes, while the other increased "
        "and was associated with extracellular defense, peptidoglycan recognition, transport, and Toll/Imd "
        "signaling. Infection therefore involved both immune activation and broad reallocation away from structural "
        "and proliferative processes. The higher mean silhouette width of the two-cluster cut supports this broad "
        "division, while the gene-level heatmap also makes clear that individual genes and locusts varied within "
        "each programme."
    )
    replace_text(
        p[127],
        "The immune-investment hypothesis predicted stronger or broader immune transcription under the "
        "carbohydrate-biased diet. Instead, recognition, signaling, and humoral-defense candidates were commonly "
        "induced within every diet, and cluster enrichment identified extracellular defense and Toll/Imd-associated "
        "functions as general features of infection. The curated list was assembled post hoc from genes significant "
        "in at least one contrast, so its category proportions are descriptive rather than independent tests. Even "
        "with that limitation, the observed pattern does not support generalized enhancement of immune transcription "
        "as the explanation for improved survival under carbohydrate-biased nutrition."
    )
    replace_text(
        p[128],
        "Earlier physiological and metabolomic work in this system showed that higher-protein diets increased "
        "hemolymph protein and branched-chain amino-acid availability and were associated with accelerated mortality "
        "and fungal development, despite elevated phenoloxidase activity (Tahir et al., 2026; Tahir et al., "
        "submitted). The present mixed responses in nutrient sensing, amino-acid metabolism, translation, and "
        "storage genes are compatible with infection-driven metabolic reorganization, but they do not establish "
        "that diet 50 or diet 83 produced a uniformly more anabolic fat body. Together, the physiological and "
        "transcriptional results are more consistent with selective differences in host metabolic state, resource "
        "availability, or tolerance than with simple immune failure, while leaving the causal mechanism unresolved."
    )
    replace_text(
        p[129],
        "Fungal RNA provided contextual evidence of among-individual heterogeneity. Most infected fat-body "
        "libraries contained little uniquely aligned fungal signal, whereas several individuals contained markedly "
        "higher values. This heterogeneity could reflect differences in infection progression, fungal localization, "
        "or sampling of a tissue in which fungal biomass was low at 96 h. The competitive reference used ARSEF 23 "
        "rather than the challenge isolate DWR2009/ARSEF 10343, and no matching tissue remained for qPCR or culture. "
        "Consequently, these alignments were not interpreted as validated fungal load, used to reclassify infected "
        "individuals, or assumed to explain host-expression variability without a formal association test."
    )
    replace_text(
        p[130],
        "The phase-reference comparison addressed whether 96 h of individual housing broadly recapitulated known "
        "density-associated expression. Of 2,594 fat-body infection DEGs in the set comparison, 2,134 did not "
        "overlap either the head or thorax phase-reference set. The smaller overlap may identify genes involved in "
        "cross-tissue physiological plasticity or shared stress responses, but the reference data came from different "
        "tissues and experimental contexts. It therefore does not demonstrate a phase transition in the experimental "
        "fat body or attribute the infection response to solitarization."
    )
    replace_text(
        p[131],
        "Sequence placement was another important source of uncertainty. A large historical DEG signal occurred on "
        "unplaced scaffolds, especially in the infected diet-33-versus-diet-50 contrast. Taxonomic, cross-genome, "
        "and homology audits did not justify treating every unplaced locus as contamination, but their provenance "
        "was less secure than that of genes on assembled chromosomes. Restricting the primary analysis to main "
        "chromosomes provides a conservative host interpretation, while preserving all-scaffold results for "
        "sensitivity and future investigation rather than deleting those loci. The major interpretation of broad "
        "immune induction and heterogeneous protein-anabolism responses was retained when all scaffolds were shown, "
        "although the curated denominator increased from 99 to 163 genes."
    )
    replace_text(
        p[132],
        "Functional enrichment in a non-model insect remains constrained by annotation coverage and orthology. Only "
        "a subset of genes in either expression cluster carried GO or KEGG assignments, and broad KEGG pathway names "
        "may reflect conserved orthologous functions rather than a locust-specific pathway definition. The use of "
        "transcript-plus-exon counts increased sensitivity relative to exon-only counting and can include signal "
        "consistent with nascent or incompletely processed RNA. The principal infection response remained extensive "
        "under exon-only counting, but the absence of an exon-only omnibus interaction makes the diet-dependent "
        "component less robust than the shared infection response. Count-definition, sequence-placement, and sample-"
        "retention results are therefore reported explicitly, and enrichment remains hypothesis-generating until "
        "supported by protein, metabolite, or functional assays."
    )
    replace_text(
        p[133],
        "Together, these results support a model in which M. robertsii induces a conserved and extensive fat body "
        "response across nutritional environments, while dietary macronutrient balance modifies a comparatively "
        "small and method-sensitive subset of responses. Canonical immune functions were induced across diets, "
        "whereas metabolic, translation, nutrient-sensing, and storage-related genes responded heterogeneously. "
        "The carbohydrate-biased diet's survival advantage is therefore unlikely to reflect generalized enhancement "
        "of immune transcription at 96 h. Future experiments combining matched "
        "fungal-load qPCR, proteomics, metabolite flux, and pathway-specific perturbation will be required to resolve "
        "whether host resource allocation, pathogen nutrition, or tolerance drives the observed survival differences."
    )

    # Replace old supplementary figure placeholders with a focused sensitivity package.
    replace_text(
        p[218],
        "This supplementary information reports sample-retention, sequence-placement, count-definition, fungal-RNA, "
        "curated-target, and formal interaction sensitivities supporting the primary all-sample, transcript-plus-"
        "exon, main-chromosome analysis."
    )
    cover_heading = p[236]
    for old_paragraph in list(doc.paragraphs[219:236]):
        remove_paragraph(old_paragraph)

    supplement_captions = [
        (
            "figureS1", 6.1, "Figure S1.",
            "Sample-retention sensitivity analysis. Primary DEG totals from all 44 libraries are compared with a "
            "secondary analysis excluding samples 1007, 1036, 1037, and 1039. This analysis evaluates the influence "
            "of atypical libraries; no sample was excluded from the primary results."
        ),
        (
            "figureS2", 6.15, "Figure S2.",
            "Chromosomal localization of differential expression. Main-chromosome DEGs are shown by genomic "
            "position and contrast context, providing a genome-wide view of whether signals are broadly distributed "
            "or concentrated in particular chromosome regions."
        ),
        (
            "figureS3", 6.15, "Figure S3.",
            "Count-definition sensitivity. Numbers of significant genes obtained with exon-only counts are compared "
            "with counts assigned from transcript and exon features and grouped by gene_id. Both matrices were "
            "depleted of annotated rRNA and restricted to the same primary chromosome scope."
        ),
        (
            "figureS4", 6.15, "Figure S4.",
            "Competitive-alignment evidence for fungal RNA. Full-range and zoomed views show uniquely aligned "
            "M. robertsii fragments and their percentage of unique competitive-reference alignments. Horizontal "
            "control means provide a background reference. These values are RNA-seq proxies, not qPCR-validated "
            "fungal loads."
        ),
        (
            "figureS5", 6.15, "Figure S5.",
            "Sequence-placement sensitivity for the curated candidate analysis. Results from all annotated scaffolds "
            "are compared with the primary main-chromosome subset, showing which changes arise from excluding "
            "unplaced or otherwise non-primary assembly sequences."
        ),
        (
            "figureS6", 6.0, "Figure S6.",
            "Formal diet-by-infection interaction sensitivity. Omnibus likelihood-ratio tests compare the full "
            "diet-plus-infection-plus-interaction model with the reduced additive model under exon-only and "
            "transcript-plus-exon counting. All 44 libraries are retained and annotated rRNA genes are excluded."
        ),
    ]

    for key, width, label, caption in supplement_captions:
        clean_name = f"{key}.png"
        copied = SUPPLEMENT_DIR / clean_name
        shutil.copy2(FIGURES[key], copied)
        add_supplement_before(cover_heading, copied, width, label, caption)

    # Bring the cover-letter summary into agreement with the revised manuscript.
    replace_text(
        p[245],
        "Our study complements earlier physiological measurements by testing the interactive effects of diet and "
        "Metarhizium infection on fat body transcription. Infection produced a large conserved response, with a "
        "467-gene core shared among all diets. Although within-diet DEG totals differed, the omnibus interaction "
        "identified only nine genes under transcript-plus-exon counting and none under exon-only counting. Immune "
        "functions were broadly induced, whereas metabolic and protein-anabolism candidates responded "
        "heterogeneously. Together with the earlier hemolymph and survival data, these results argue against "
        "generalized enhancement of immune transcription as the carbohydrate-associated survival mechanism and "
        "instead motivate tests of host resource allocation, pathogen nutrition, and disease tolerance."
    )

    # Keep figures and captions together where possible and use journal-friendly body spacing.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.lower() in {"normal", "body text"}:
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.08

    doc.core_properties.title = (
        "Fat body transcriptional responses to dietary macronutrient balance and Metarhizium infection"
    )
    doc.core_properties.comments = (
        "Clean CBP Part B revision generated from current all-sample, transcript-plus-exon, rRNA-depleted, "
        "main-chromosome analysis outputs."
    )
    doc.save(OUTPUT_DOCX)

    # A separate review copy highlights revised passages without marking the clean manuscript.
    review = Document(OUTPUT_DOCX)
    review_paragraphs = review.paragraphs
    note = review_paragraphs[0].insert_paragraph_before(
        "REVIEW COPY: passages revised from the supplied manuscript are highlighted in yellow."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].bold = True
    note.runs[0].font.size = Pt(10)

    highlight = False
    highlight_single_starts = (
        "Fungal infection strongly reshaped",
        "Infection-responsive expression was broadly shared",
        "A core of 467 infection-responsive",
        "Formal interaction tests supported selective",
        "Dietary macronutrient balance alters disease outcome",
        "This supplementary information reports sample-retention",
        "Our study complements earlier physiological measurements",
    )
    for paragraph in review.paragraphs:
        text = paragraph.text.strip()
        if text in {"2. Materials and methods", "3. Results", "4. Discussion"}:
            highlight = True
        if text in {"Acknowledgements", "Draft cover letter"}:
            highlight = False
        should_highlight = highlight or text.startswith(highlight_single_starts)
        if should_highlight:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Highlight the replacement supplementary figure package up to the cover letter.
    in_supplements = False
    for paragraph in review.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("This supplementary information reports sample-retention"):
            in_supplements = True
        if text == "Draft cover letter":
            in_supplements = False
        if in_supplements:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    review.core_properties.comments = (
        "Review copy of the CBP Part B revision. Yellow highlighting identifies rewritten passages."
    )
    review.save(OUTPUT_REVIEW_DOCX)
    print(OUTPUT_DOCX)
    print(OUTPUT_REVIEW_DOCX)


if __name__ == "__main__":
    main()
