#!/usr/bin/env python3
"""Build a focused Word audit of the latest Materials and Methods and Results.

The source manuscript is never modified. The output contains proposed replacement
text based on the final 44-library, main-chromosome, transcript-plus-exon analysis.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[2]
SOURCE = Path(
    "/Users/maevatecher/Desktop/"
    "Chapter # 3 Understanding the molecular mechanisms associated with diet and immune function (2).docx"
)
OUT_DIR = PROJECT / "manuscript" / "revisions" / "20260810_methods_results_latest_analysis_audit"
OUT_DOCX = OUT_DIR / "Chapter_3_methods_results_latest_analysis_REVIEW_20260810.docx"


# compact_reference_guide preset, resolved here rather than relying on Word defaults.
INK = "172B4D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5C6675"
LIGHT_BLUE = "E8EEF5"
LIGHT_YELLOW = "FFF7D6"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
RISK_RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color: str, side="left", size=16, space=8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False

    header = section.header.paragraphs[0]
    header.text = "Chapter 3 | Methods and Results revision audit"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    set_paragraph_border(header, "D9E1EA", side="bottom", size=6, space=4)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Chapter 3: latest-analysis revision audit")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Proposed replacement text for Materials and methods, Results, and figure legends")
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    callout = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(callout, [9360])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run("Purpose. ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(
        "This review copy shows the sections that should be replaced before the next manuscript audit. "
        "The source manuscript remains unchanged. Yellow blocks are ready-to-paste proposed text; blue notes "
        "identify points that still require author confirmation or a deliberate reporting choice."
    )


def add_proposed_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_after = Pt(7)
    shade_paragraph(p, LIGHT_YELLOW)
    set_paragraph_border(p, "D6A700", side="left", size=12, space=6)


def add_note(doc: Document, title: str, text: str, risk=False) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FCEBEC" if risk else LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(title + " ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(RISK_RED if risk else DARK_BLUE)
    p.add_run(text)


def add_change_table(doc: Document) -> None:
    doc.add_heading("High-priority corrections", level=1)
    rows = [
        ("Cohort", "45 libraries were sequenced; 1044 failed host-alignment QC; 44 libraries were analyzed. Corrected sample 1024 is infected diet 33."),
        ("Group sizes", "Diet 33: 7 control, 10 infected; diet 50: 6 control, 9 infected; diet 83: 7 control, 5 infected."),
        ("rRNA", "The input matrix contained 95,606 rows. The value 12,402 is the number of rRNA identifiers removed, not the number of genes in the matrix."),
        ("Primary counts", "Competitive host-pathogen alignment; featureCounts -t transcript,exon -g gene_id; primary, high-quality paired fragments; main nuclear chromosomes."),
        ("Outliers", "No PCA-defined sample was removed after QC. The earlier 1007/1037 removal exercise is not part of the primary analysis."),
        ("PCA", "PC1 = 27.7%; PC2 = 12.4%, not 32.7% and 21.0%."),
        ("Within-diet DEGs", "Diet 33 = 1,036; diet 50 = 2,000; diet 83 = 1,013."),
        ("Infection overlap", "2,581 unique infection-responsive genes; 472 shared by all diets; 996 shared by at least two diets."),
        ("Expression clusters", "2,918 main-chromosome DEGs separated into cluster 1 (1,237) and cluster 2 (1,681). GO/KEGG panels use only within-diet infection DEGs in each cluster."),
        ("Curated candidates", "All 111 curated genes are on main chromosomes. Status tests are exploratory because curation was informed partly by earlier analyses of the same data."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_fixed_width(table, [2700, 6660])
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Draft area"
    hdr[1].text = "Required correction"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(INK)
    for label, detail in rows:
        row = table.add_row()
        keep_table_row_together(row)
        cells = row.cells
        cells[0].text = label
        cells[1].text = detail
        cells[0].paragraphs[0].runs[0].bold = True
        if len(table.rows) % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)


METHODS = [
    (
        "2.1 Experimental design",
        [
            "We used a 3 x 2 factorial design comprising three isocaloric diets and two inoculation treatments. The diets contained 33%, 50%, or 83% protein as a proportion of total dietary macronutrient, corresponding to whole-diet protein:carbohydrate percentages of 14:28, 21:21, and 35:7, respectively. Individual fifth-instar female Schistocerca gregaria were randomly assigned to one diet and received either Metarhizium robertsii conidia suspended in oil or sterile oil alone.",
            "Forty-five RNA-seq libraries were generated. Library 1044 (infected, diet 83) was excluded before downstream host-expression analysis because it failed host-alignment quality control, with 1.78% uniquely mapped reads and 3.23% uniquely plus multimapped reads. The final analysis therefore included 44 libraries: diet 33 comprised seven controls and 10 infected locusts, diet 50 comprised six controls and nine infected locusts, and diet 83 comprised seven controls and five infected locusts (20 controls and 24 infected locusts in total). Sample 1024 was assigned to infected diet 33 according to the original experimental datasheet. No additional library was removed on the basis of its position in exploratory PCA or sample-distance analyses."
        ],
    ),
    (
        "2.2 Host, pathogen, and diets",
        [
            "Fifth (final) instar female S. gregaria, three days post-molt, were obtained from a gregarious colony maintained at Arizona State University from laboratory lines originating at the University of Leicester (United Kingdom) and the University of Haifa-Oranim (Israel). Locusts were maintained at 34 degrees C during the day and 25 degrees C at night, at 20-50% relative humidity under a 14 h:10 h light:dark cycle. During the experiment, locusts were housed individually in ventilated plastic containers (18.9 x 13.5 x 9.5 cm) and received food and water ad libitum.",
            "Locusts were challenged with M. robertsii isolate DWR2009, also designated ARSEF 10343, which was originally isolated from soil in the United States. Conidia were produced by biphasic solid-state fermentation, suspended in Orchex 792 paraffinic oil, and adjusted to 1 x 10^9 conidia mL-1. Concentration was determined with a hemocytometer and viability was confirmed by germination on potato dextrose agar at 25-26 degrees C. Infected locusts received 1 microliter of the conidial suspension at the hind-leg femur-trochanter junction. Controls received 1 microliter of sterile Orchex 792 oil and were handled identically.",
            "Chemically defined diets followed the nutritional framework of Dadd (1961) and Simpson and Abisgold (1985), with supplements as described by Zembrzuski et al. (2023). Each diet contained 42% macronutrient by dry mass. Protein was supplied as a 3:1:1 mixture of casein, peptone, and albumen, whereas carbohydrate was supplied as equal proportions of sucrose and dextrin. Diet and water were replenished as needed."
        ],
    ),
    (
        "2.3 Tissue sampling, RNA extraction, and sequencing",
        [
            "Abdominal fat body was collected 96 h after inoculation. A single experimenter performed dissections under sterile, RNase-free conditions. Instruments were flame sterilized and sequentially treated with 10% bleach, distilled water, 70% ethanol, and RNaseZap before each dissection. Locusts were surface sterilized with 70% ethanol, and fat body was removed while avoiding the gut, trachea, and heart. Tissue was transferred immediately to nuclease-free tubes, flash-frozen in liquid nitrogen, and stored at -80 degrees C.",
            "Total RNA was extracted on a Promega Maxwell RSC instrument with the simplyRNA Tissue Kit according to the manufacturer's protocol. RNA concentration was measured with a Qubit 4 Fluorometer and the Qubit RNA HS Assay. Stranded total-RNA libraries were prepared with the Illumina Stranded Total RNA kit and Ribo-Zero rRNA depletion. Pooled libraries were sequenced by Novogene on an Illumina NovaSeq X Plus 25B platform as 150-bp paired-end reads, targeting approximately 45 million read pairs per library."
        ],
    ),
    (
        "2.4 Read processing, competitive mapping, and taxonomic quality control",
        [
            "RNA-seq preprocessing was executed on the Arizona State University high-performance computing cluster with a run-scoped Snakemake workflow. Paired reads were processed with fastp using automatic paired-end adapter detection, removal of the first two bases of each read, and a minimum retained length of 50 bp. Per-library JSON and HTML quality reports were retained.",
            "Trimmed reads were aligned in two-pass mode with STAR v2.7.10b. The primary branch used a competitive reference comprising the S. gregaria assembly GCF_023897955.1 and the M. robertsii assembly GCF_000187425.2. The fungal assembly represents ARSEF 23 and was used as a close mapping proxy rather than a strain-matched reference for the experimental DWR2009/ARSEF 10343 isolate. STAR indices were constructed with sjdbOverhang = 149 and alignIntronMax = 2,500,000; up to 20 mapping locations were reported and unique alignments were assigned MAPQ 60. A host-only STAR branch was retained as a mapping sensitivity analysis.",
            "Fungal-alignment burden was summarized from uniquely aligned fragments assigned to the fungal component of the competitive reference and expressed as a percentage of uniquely aligned competitive-reference fragments. Kraken2 v2.17.1 with the PlusPF 20260626 database and Bracken family-level estimates were used as complementary taxonomic screens of trimmed and host-unmapped reads. Because the fungal reference was not strain matched and no tissue remained for qPCR or culture-based confirmation, RNA-seq fungal alignment was interpreted only as evidence of fungal transcriptional signal, not as a validated measure of pathogen load."
        ],
    ),
    (
        "2.5 Gene counting, rRNA removal, and sequence placement",
        [
            "Gene-level fragments were quantified with featureCounts from the Subread package using paired-end fragment counting (-p --countReadPairs), properly paired and chimeric-fragment checks (-B -C), reverse-strand specificity (-s 2), primary alignments (--primary), a minimum mapping quality of 10 (-Q 10), transcript and exon feature types (-t transcript,exon), and gene_id as the grouping attribute (-g gene_id). Multimapping fragments were not counted because the -M option was not used and their STAR mapping qualities fell below the counting threshold. A conventional exon-only matrix was processed in parallel as a count-definition sensitivity analysis.",
            "The unfiltered count matrix contained 95,606 gene identifiers. All 12,402 identifiers in the project rRNA exclusion list were removed, leaving 83,204 non-rRNA identifiers before expression and placement filters. Residual rRNA accounted for a median 7.87% of assigned transcript-plus-exon fragments per retained library (range 3.28-41.70%). For the primary analysis, sequence placement was assigned from the official NCBI assembly report and analysis was restricted to non-rRNA genes on the nuclear chromosome pseudomolecules. All-scaffold results were retained as a placement sensitivity analysis."
        ],
    ),
    (
        "2.6 Differential expression and multivariate visualization",
        [
            "Differential expression was analyzed from integer counts with DESeq2. Genes were retained when they had at least 10 counts in at least three libraries. The primary transcript-plus-exon, main-chromosome analysis included 14,132 expressed genes. DESeq2 estimated sample-specific size factors to account for differences in sequencing depth and library composition during model fitting. Automatic count replacement was disabled (minReplicatesForReplace = Inf), and Cook's-distance filtering was disabled for planned contrasts (cooksCutoff = FALSE), so all 44 retained libraries contributed to model estimates. Differential expression required a Benjamini-Hochberg adjusted P value below 0.05 and an absolute log2 fold change of at least 1. Variance-stabilized values derived from the normalized counts were used only for PCA and heatmaps, not as input to differential-expression testing.",
            "A six-level diet-by-treatment no-intercept model (~ 0 + group) generated exact within-diet infected-versus-control contrasts and pairwise diet contrasts separately among control and infected locusts. A model containing diet and treatment main effects (~ Diet + Treatment) estimated the average infection response across diets. A full factorial model (~ Diet + Treatment + Diet:Treatment) and a likelihood-ratio comparison against ~ Diet + Treatment assessed whether the infection response differed among diets. For diet contrasts, diet 50 served as the denominator for comparisons with diets 33 and 83, and diet 33 served as the denominator for the 83-versus-33 comparison. Positive log2 fold changes indicate higher expression in the named numerator condition.",
            "PCA was calculated from variance-stabilized expression of the rRNA-filtered primary gene set. Group centroids were calculated for each diet-by-treatment combination, and lines joining control and infected centroids summarized the direction and magnitude of the infection-associated shift within each diet. No library was removed because it occupied an extreme position in the PCA."
        ],
    ),
    (
        "2.7 DEG overlap, assessment of phase change as a potential confounder, expression clustering, and functional enrichment",
        [
            "Overlap among the three within-diet infection DEG sets was summarized by set intersection. Because the infection protocol required locusts to be maintained individually under semi-isolated conditions for 96 h, we assessed whether early density-dependent phase change could have contributed to the heterogeneous fat body expression response and therefore acted as a potential confounding factor. The union of the within-diet infection DEGs was compared with independent solitarious-versus-gregarious DEG sets from S. gregaria head and thorax (Techer et al., in prep.). The reference study also used female final-instar nymphs sampled three days post-molt and the same total-RNA, rRNA-depletion library-preparation workflow, reducing differences attributable to sex, developmental stage, and library preparation. Head and thorax were evaluated separately, and genes detected in both tissues were treated as the most conservative cross-tissue phase-associated signature. Overlap with only one tissue was retained as tissue-dependent phase evidence. Such overlaps identify infection-responsive genes whose expression could reflect infection, an early response to reduced social density, or both; they do not determine which process caused the response. Genes absent from both reference sets were described as lacking support in the available phase references, rather than as proven infection-specific, because the reference profiles came from different tissues and an independent experiment. The comparison therefore assessed phase change as a possible source of confounding but did not establish that a phase transition occurred in the experimental fat body samples.",
            "Separately, to identify coordinated expression programs within the fat body response, all significant main-chromosome DEGs across the displayed fat body contrasts were clustered from row-standardized variance-stabilized expression. Euclidean distances and complete-linkage hierarchical clustering were used, and the primary heatmap was divided into two major expression clusters by cutting the same gene dendrogram at k = 2. Individual-gene expression was retained; cluster averages were not substituted for gene-level values. Within each cluster, genes were ordered by DEG context and statistical evidence for display.",
            "Functional descriptions, Gene Ontology (GO) terms, and KEGG orthology and pathway assignments were obtained from the project eggNOG annotation and linked to count-matrix gene identifiers through CDS records in the reference GFF. Over-representation tests used clusterProfiler::enricher and the expressed, non-rRNA main-chromosome genes carrying the relevant annotation as the background. P values were adjusted by the Benjamini-Hochberg method. Figure 3 GO and KEGG foregrounds were restricted to genes called differentially expressed in at least one within-diet infection contrast within the corresponding expression cluster. KEGG results were interpreted as orthology-based, hypothesis-generating annotations; the eukaryote-inclusive view excluded explicitly prokaryotic and microbial disease-reference pathways, and an insect/arthropod-focused view was retained as a sensitivity analysis."
        ],
    ),
    (
        "2.8 Curated immune and protein-anabolism candidates",
        [
            "A frozen set of 111 biologically curated genes was used to examine immune and protein-anabolism responses. The set comprised 77 immune candidates (cellular immunity, humoral immunity, and innate immune recognition and signaling) and 34 protein-anabolism candidates (nitrogen and amino-acid transport and metabolism, nutrient sensing and TOR/insulin signaling, protein synthesis and translation, and storage hexamerins and reserves). All 111 candidates were located on the main chromosomes and were retained regardless of whether they met the final differential-expression threshold.",
            "For each planned contrast, candidates were classified as higher in the numerator condition, higher in the denominator condition, or not significantly different using the genome-wide DEG criteria. Chi-square tests compared these status distributions among diets; when expected cell counts were below five, P values were estimated by Monte Carlo simulation with 20,000 replicates. Because the curated set was informed partly by biological annotation and partly by earlier screens of the same experiment, these tests were treated as exploratory summaries rather than independent confirmatory tests."
        ],
    ),
]


RESULTS = [
    (
        "3.1 Library quality and global fat body expression",
        [
            "Of 45 sequenced libraries, 44 passed host-alignment quality control and entered the expression analysis. Library 1044 had 1.78% uniquely mapped and 1.45% multimapped reads, with 89.81% of reads classified by STAR as unmapped because they were too short; it was therefore excluded before normalization and differential-expression testing. Sample 1024 was analyzed as infected diet 33. All other libraries, including samples with comparatively dispersed PCA positions, were retained.",
            "The first two principal components of variance-stabilized expression explained 27.7% and 12.4% of total variance, respectively (Fig. 1). Control samples were concentrated toward negative PC1 values, whereas most infected samples shifted toward positive PC1 values. Control-to-infected centroid displacement was observed in all diets and was largest in diet 50, followed by diets 83 and 33. The broad and partly overlapping infected polygons nevertheless showed substantial inter-individual variation, particularly in diets 33 and 83.",
            "Fungal alignment remained near background in most fat body libraries but was highly heterogeneous among infected individuals. Median uniquely aligned fungal percentages were 0.00123%, 0.00097%, and 0.00060% in infected diets 33, 50, and 83, respectively, and the maximum infected-library value was 3.13%. Control medians ranged from 0.00025% to 0.00056%. These values were used as an RNA-seq fungal-signal proxy rather than a quantitative measure of pathogen load."
        ],
    ),
    (
        "3.2 Infection dominated differential expression, with diet-dependent breadth",
        [
            "The average infected-versus-control contrast across diets identified 1,696 DEGs, including 907 genes with higher expression in infected locusts and 789 with higher expression in controls (Fig. 2A). Within diets, infection altered 1,036 genes in diet 33 (577 higher in infected and 459 higher in controls), 2,000 genes in diet 50 (1,260 and 740), and 1,013 genes in diet 83 (663 and 350). Protein-coding genes formed the majority of each response, but lncRNAs contributed 56, 132, and 55 DEGs in diets 33, 50, and 83, respectively.",
            "Direct diet contrasts were much smaller among controls, identifying 13 DEGs between diets 33 and 50, 61 between diets 83 and 33, and 57 between diets 83 and 50. Among infected locusts, the corresponding diet contrasts identified 242, 105, and 62 DEGs. Thus, the largest direct diet contrast was still substantially smaller than any within-diet infection response.",
            "The three within-diet infection contrasts contained 2,581 unique DEGs (Fig. 2B). Of these, 472 were shared by all three diets and 996 (38.6%) occurred in at least two diets. Diet-specific components comprised 210 genes unique to diet 33, 1,067 unique to diet 50, and 308 unique to diet 83. The large diet-50-specific component accounted for much of the greater infection-associated DEG burden in the balanced diet.",
            "The fat body heatmaps revealed marked inter-individual heterogeneity, including DEG groups with strong expression in only a small number of locusts. Comparison with the independent density-phase references showed that 2,117 of the 2,581 infection-responsive genes (82.0%) were not detected as phase DEGs in either head or thorax (Fig. 2C). Eighty-two infection-responsive genes overlapped the head set only, 259 overlapped the thorax set only, and 123 were also phase-responsive in both reference tissues. These 123 genes provide the strongest evidence that a cross-tissue density-dependent program could contribute to part of the apparent infection response, while the one-tissue overlaps provide weaker, tissue-dependent evidence of potential confounding. The majority of infection DEGs therefore showed no phase association in the available references. However, this absence cannot exclude phase-related confounding or establish infection specificity because fat body was not represented in the independent phase study."
        ],
    ),
    (
        "3.3 Two expression programs separated suppressed metabolic and cuticle functions from induced defense functions",
        [
            "The union of 2,918 significant main-chromosome DEGs across the displayed fat body contrasts separated into two major expression clusters (Fig. 3A). Cluster 1 contained 1,237 genes and was generally characterized by higher expression in controls and lower expression after infection. Cluster 2 contained 1,681 genes and showed the converse pattern, with broadly higher expression in infected locusts. The functional-enrichment panels focused on the 1,129 cluster-1 genes and 1,337 cluster-2 genes that were DEGs in at least one within-diet infection contrast.",
            "Cluster 1 was enriched for GO terms associated with chitin-based cuticle development, cuticle structure, extracellular matrix, cell-cycle processes, and oxidoreductase and monooxygenase activities (Fig. 3B). Its eukaryote-inclusive KEGG results included insect hormone biosynthesis, cell cycle, xenobiotic and drug metabolism by cytochrome P450, pentose and glucuronate interconversions, and DNA replication (Fig. 3C). These annotations are consistent with reduced representation of developmental, structural, and metabolic programs during infection.",
            "Cluster 2 was enriched for response to bacterium, response to biotic stimulus, defense response, peptidoglycan metabolism and binding, extracellular region, plasma membrane, signaling, and transmembrane transport. KEGG enrichment included Toll and Imd signaling and the fly Hedgehog signaling pathway. Together, the two clusters describe a broad shift from cuticle, cell-cycle, and metabolic expression toward extracellular defense, microbial recognition, signaling, and transport during infection."
        ],
    ),
    (
        "3.4 Curated immune and protein-anabolism candidates",
        [
            "Among the 77 curated immune genes, infection was associated with higher expression of 37 genes in diet 33, 54 in diet 50, and 39 in diet 83; six, 12, and one genes, respectively, were higher in controls (Fig. 4). The exploratory distribution of higher-in-infected, higher-in-control, and non-significant statuses differed among diets (Monte Carlo chi-square P < 0.0001), with diet 50 showing the largest fraction of immune candidates with higher expression during infection (70.1%). This difference was driven principally by innate immune recognition and signaling candidates. Humoral and cellular subsets were smaller and did not show statistically resolved diet differences in their status distributions.",
            "Among the 34 protein-anabolism candidates, infection was associated with higher expression of nine genes in diet 33, 17 in diet 50, and 10 in diet 83, whereas three, 11, and one genes, respectively, were higher in controls. The exploratory status distribution again differed among diets (Monte Carlo chi-square P < 0.0001), with diet 50 producing both the largest upregulated component and the largest downregulated component. Nitrogen and amino-acid transport and metabolism accounted for much of this pattern; the smaller storage, protein-synthesis, and TOR/insulin categories were less conclusive.",
            "Direct diet contrasts among infected locusts produced few changes within the curated sets. For immune candidates, 74-75 of 77 genes were not significant in each pairwise diet comparison, and the status distribution did not differ among contrasts. For protein-anabolism candidates, four genes were higher in diet 50 than diet 33, one was lower in diet 83 than diet 33, and none differed between diets 83 and 50. These sparse direct diet effects do not support broad diet-specific activation of the curated pathways after infection, despite the larger within-diet infection response observed in diet 50."
        ],
    ),
    (
        "3.5 Sensitivity to count definition and sequence placement",
        [
            "The transcript-plus-exon primary definition retained 14,132 expressed main-chromosome genes, whereas exon-only counting retained 11,140. Within-diet infection calls overlapped substantially but were more numerous under transcript-plus-exon counting: diet 33 contained 614 shared DEGs, 219 exon-only DEGs, and 422 transcript-plus-exon-only DEGs; diet 50 contained 891 shared, 222 exon-only, and 1,109 transcript-plus-exon-only DEGs; and diet 83 contained 556 shared, 202 exon-only, and 457 transcript-plus-exon-only DEGs. The qualitative conclusion that infection dominated the response and that diet 50 had the broadest response was retained, but absolute DEG counts were count-definition dependent.",
            "The formal omnibus diet-by-infection interaction was also count-definition sensitive: no exon-only gene met the adjusted-P threshold, whereas 17 transcript-plus-exon genes met the adjusted-P threshold and the pairwise log2-fold-change criterion on the main chromosomes. All-scaffold analyses were retained as a separate placement sensitivity. Because unresolved unplaced scaffolds disproportionately affected some contrasts, particularly infected diet comparisons, main-chromosome results were used for the primary biological interpretation."
        ],
    ),
]


LEGENDS = [
    (
        "Figure 1",
        "Principal component analysis of the final fat body RNA-seq cohort. PCA used variance-stabilized counts from expressed, non-rRNA genes on the main S. gregaria chromosomes. Points represent individual libraries, shapes indicate treatment, and colors indicate diet. Shaded polygons show the observed breadth of each diet-by-treatment group, open circles mark group centroids, and lines connect control and infected centroids within each diet."
    ),
    (
        "Figure 2",
        "Differential-expression burden and overlap in the transcript-plus-exon, main-chromosome analysis. (A) Numbers and directions of significant protein-coding and long non-coding RNA DEGs for the average infection response, infection within each diet, diet contrasts among controls, and diet contrasts among infected locusts. Red denotes higher expression in the named numerator condition and blue denotes higher expression in the denominator. (B) Overlap among within-diet infection DEG sets. (C) Overlap between the union of infection-responsive fat body DEGs and independent head and thorax solitarious-versus-gregarious DEG sets. External overlap provides phase-related context but does not establish a phase effect in fat body. Annotated rRNA genes were absent from all sets."
    ),
    (
        "Figure 3",
        "Fungal RNA signal and two major expression programs among main-chromosome DEGs. (A) The upper track shows uniquely aligned fungal reads as a percentage of unique competitive-reference alignments; values above 0.1% are capped for display and labeled with their full value. The heatmap shows row-standardized variance-stabilized expression for individual genes, with samples grouped by treatment and diet. Genes are divided by a complete-linkage dendrogram cut at k = 2. (B) Significant GO enrichment for infection-within-diet DEGs in each cluster. (C) Significant eukaryote-inclusive KEGG enrichment for the same cluster-specific foregrounds. Dot size represents the number of foreground genes assigned to each term."
    ),
    (
        "Figure 4",
        "Expression status of 111 curated immune and protein-anabolism genes. Within-diet panels classify infected-versus-control responses, whereas among-infected panels classify pairwise diet contrasts. Bars show the percentage of curated targets with higher expression in the numerator condition, higher expression in the denominator, or no significant difference at adjusted P < 0.05 and absolute log2 fold change >= 1. Numbers below categories indicate the number of curated genes tested. Chi-square P values are exploratory because the candidate list was informed partly by earlier analyses of the same experiment."
    ),
]


SUPPLEMENT_ROWS = [
    ("Figure S1", "Library mapping QC and exclusion of 1044", "Stack unique mapping above unique plus multimapping; label the diet x treatment model tests and show 1044 explicitly."),
    ("Figure S2", "Count-definition sensitivity", "Compare transcript-plus-exon and exon-only expressed genes, DEG totals, overlaps, and the omnibus interaction result."),
    ("Figure S3", "Sequence-placement sensitivity", "Compare main-chromosome and all-scaffold DEG burdens and summarize the unplaced-scaffold homology audit."),
    ("Figure S4", "Fungal RNA and taxonomic evidence", "Show competitive fungal-alignment burden together with Kraken2/Bracken family-level summaries; retain the caveat that this is not qPCR-confirmed load."),
    ("Figure S5", "Body mass and mapping quality", "Show mass against unique and unique-plus-multimapping rates with diet and infection annotations and adjusted statistical tests."),
    ("Table S1", "Sample metadata and QC", "Include corrected 1024 assignment, 1044 exclusion, group membership, input reads, STAR metrics, and fungal-alignment percentage."),
    ("Table S2", "Complete primary DEG catalogue", "Report every tested gene and contrast, log2 fold change, adjusted P value, direction, biotype, placement, and GFF description."),
    ("Table S3", "DEG intersection membership", "List infection-within-diet and phase-reference set membership for every gene used in Figure 2."),
    ("Table S4", "GO and KEGG reverse index", "List each enriched term, adjusted P value, foreground/background counts, and all contributing gene identifiers."),
    ("Table S5", "Curated candidate audit", "Provide all 111 genes, functional categories, placement, descriptions, and status in every displayed contrast."),
]


def add_section_blocks(doc: Document, heading: str, blocks) -> None:
    doc.add_heading(heading, level=1)
    for subsection, paragraphs in blocks:
        doc.add_heading(subsection, level=2)
        for text in paragraphs:
            add_proposed_paragraph(doc, text)


def add_legends(doc: Document) -> None:
    doc.add_heading("Proposed updated figure legends", level=1)
    for label, text in LEGENDS:
        doc.add_heading(label, level=2)
        add_proposed_paragraph(doc, text)


def add_supplement(doc: Document) -> None:
    doc.add_heading("Recommended supplementary package", level=1)
    p = doc.add_paragraph(
        "These items support the main interpretation without reintroducing archived PCA-defined outlier-removal analyses into the primary narrative."
    )
    p.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=1, cols=3)
    set_table_fixed_width(table, [1500, 3000, 4860])
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("Item", "Purpose", "Minimum content")):
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        cell.paragraphs[0].runs[0].bold = True
    set_repeat_table_header(table.rows[0])
    for item, purpose, content in SUPPLEMENT_ROWS:
        row = table.add_row()
        keep_table_row_together(row)
        cells = row.cells
        cells[0].text = item
        cells[1].text = purpose
        cells[2].text = content
        cells[0].paragraphs[0].runs[0].bold = True
        if len(table.rows) % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)


def add_author_checks(doc: Document) -> None:
    doc.add_heading("Author checks before submission", level=1)
    checks = [
        ("Library preparation site", "Confirm whether libraries were prepared at Texas A&M University or in-house, and retain only the verified wording."),
        ("RNA quality criterion", "Insert the instrument and numerical RNA-integrity acceptance criterion if one was applied."),
        ("Software versions", "Add fastp, Subread/featureCounts, DESeq2, clusterProfiler, and eggNOG-mapper versions from the frozen run provenance."),
        ("Fungal isolate wording", "Keep DWR2009/ARSEF 10343 as the challenge isolate and ARSEF 23/GCF_000187425.2 as the mapping proxy; do not merge these identities."),
        ("Curated-gene inference", "Decide whether exploratory chi-square P values remain in the main text. The status counts are useful, but the tests are not independent of candidate selection."),
        ("Discussion numbers", "Update every DEG, overlap, PCA, and curated-candidate value repeated in the Discussion after the Results wording is accepted."),
    ]
    for title, text in checks:
        add_note(doc, title + ".", text, risk=title in {"Library preparation site", "Curated-gene inference"})


def main() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(f"Missing source manuscript: {SOURCE}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_document(doc)
    add_title(doc)
    add_change_table(doc)

    add_section_blocks(doc, "Proposed replacement: Materials and methods", METHODS)
    add_note(
        doc,
        "Reporting choice.",
        "The primary manuscript analysis is transcript plus exon, non-rRNA, competitive-host, and main-chromosome. Exon-only, host-only, and all-scaffold branches are sensitivities and should not be blended into the primary DEG totals."
    )

    add_section_blocks(doc, "Proposed replacement: Results", RESULTS)

    add_legends(doc)
    add_supplement(doc)
    add_author_checks(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Source reviewed: {SOURCE.name}")
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
