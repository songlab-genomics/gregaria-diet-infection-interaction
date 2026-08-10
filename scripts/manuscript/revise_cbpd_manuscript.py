#!/usr/bin/env python3
"""Create a clean CBP Part D manuscript revision from the current Word draft.

The script never edits the source manuscript. It writes a dated revision, embeds
the current publication figures, and copies the proposed supplementary figures
to the same revision folder so the provenance of every panel remains explicit.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


PROJECT = Path(__file__).resolve().parents[2]
SOURCE = Path(
    "/Users/maevatecher/Desktop/"
    "Chapter # 3 Understanding the molecular mechanisms associated with diet and immune function (1).docx"
)
REVISION_DIR = PROJECT / "manuscript" / "revisions" / "20260807_cbpd_methods_results_update"
OUTPUT_DOCX = REVISION_DIR / "Chapter_3_CBPD_methods_results_revised_20260807.docx"
OUTPUT_REVIEW_DOCX = REVISION_DIR / "Chapter_3_CBPD_revised_changes_highlighted_20260807.docx"
SUPPLEMENT_DIR = REVISION_DIR / "supplementary_figures"

RUN_PUBLICATION = PROJECT / "output" / "rmd_runs" / "publication-figures_20260807-171511"
RUN_COUNT_SENSITIVITY = PROJECT / "output" / "rmd_runs" / "count-definition-sensitivity_20260807-155445"
RUN_FUNGAL = PROJECT / "output" / "rmd_runs" / "metarhizium-read-burden_20260807-143116"
RUN_QC = PROJECT / "output" / "rmd_runs" / "count-qc-exploration_20260810-144845"

FIGURES = {
    "figure1": RUN_PUBLICATION / "figure1_fatbody_pca_centroid_shifts.png",
    "figure2": RUN_PUBLICATION / "figure2_deg_burden_context_phase_venns.png",
    "figure3": RUN_PUBLICATION / "figure3_expression_cluster_gene_level_GO_KEGG_lollipop_eukaryote_inclusive.png",
    "figure4": RUN_PUBLICATION / "figure4_curated_targets_main_chromosomes.png",
    "figureS1": RUN_QC / "figureS3_mapping_quality_exclusion_1044.png",
    "figureS2": RUN_PUBLICATION / "figureS2_deg_chromosome_localization.png",
    "figureS3": RUN_COUNT_SENSITIVITY / "count_definition_deg_totals.png",
    "figureS4": RUN_FUNGAL / "metarhizium_alignment_burden_combined.png",
    "figureS5": RUN_PUBLICATION / "figure4_curated_target_placement_impact.png",
}


def clear_paragraph(paragraph):
    """Remove text and drawings while preserving the paragraph properties."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_text(paragraph, text: str, style: str | None = None, size: float | None = None):
    clear_paragraph(paragraph)
    if style is not None:
        paragraph.style = style
    run = paragraph.add_run(text)
    if size is not None:
        run.font.size = Pt(size)
    return paragraph


def blank(paragraph):
    clear_paragraph(paragraph)


def add_caption(paragraph, label: str, text: str):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(8)
    lead = paragraph.add_run(label + " ")
    lead.bold = True
    lead.font.size = Pt(9.5)
    body = paragraph.add_run(text)
    body.font.size = Pt(9.5)


def add_picture(paragraph, path: Path, width: float):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def add_supplement_before(target, image: Path, width: float, label: str, caption: str):
    image_paragraph = target.insert_paragraph_before()
    image_paragraph.paragraph_format.page_break_before = True
    add_picture(image_paragraph, image, width)
    caption_paragraph = target.insert_paragraph_before()
    add_caption(caption_paragraph, label, caption)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def validate_inputs():
    missing = [str(path) for path in [SOURCE, *FIGURES.values()] if not path.is_file()]
    if missing:
        raise FileNotFoundError("Missing manuscript input(s):\n" + "\n".join(missing))


def main():
    validate_inputs()
    REVISION_DIR.mkdir(parents=True, exist_ok=True)
    SUPPLEMENT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(SOURCE)
    p = doc.paragraphs

    # Correct the author display while preserving the existing superscript affiliations.
    for author_paragraph in (p[5], p[207]):
        for run in author_paragraph.runs:
            run.text = run.text.replace("Maeva Techer", "Maeva A. Techer")

    # Submission highlights: each statement remains under 100 characters.
    replace_text(p[27], "Fungal infection strongly reshaped the desert locust fat body transcriptome")
    replace_text(p[28], "Diet 50 produced the broadest transcriptional response to fungal infection")
    replace_text(p[29], "A core of 467 infection-responsive genes was shared among all three diets")
    replace_text(p[30], "Curated immune and anabolic responses did not differ significantly among diets")

    abstract = (
        "Dietary macronutrient balance can alter disease outcome, but the molecular mechanisms linking "
        "nutrition to host defense remain unresolved. We tested how dietary protein-to-carbohydrate ratio "
        "modifies the fat body transcriptional response of gregarious desert locusts (Schistocerca gregaria) "
        "to infection by the entomopathogenic fungus Metarhizium robertsii. Forty-four fifth-instar females "
        "were assigned to carbohydrate-biased (diet 33), balanced (diet 50), or protein-biased (diet 83) "
        "diets and infected or oil-control treatments. All 44 samples were retained. Competitive mapping and "
        "transcript-plus-exon counting preceded rRNA removal and main-chromosome differential-expression analysis. "
        "Infection altered 1,016, 2,009, and 1,030 genes within "
        "diets 33, 50, and 83, respectively, whereas direct diet contrasts identified 15-60 genes among "
        "controls and 59-303 genes among infected locusts. The three within-diet infection responses shared "
        "467 genes. Two broad expression programs separated genes that decreased during infection and were "
        "enriched for cuticle, cell-cycle, and metabolic functions from genes that increased during infection "
        "and were enriched for extracellular defense, peptidoglycan recognition, transport, and Toll/Imd "
        "signaling. Curated immune and protein-anabolism gene categories showed diet-associated tendencies, "
        "but their proportions of upregulated, downregulated, and non-significant genes did not differ "
        "significantly among diets. Fungal-alignment signal was highly heterogeneous among infected fat-body "
        "libraries and was treated as an RNA-seq proxy rather than a validated measure of fungal load. These "
        "results indicate that infection dominates fat body transcription, while diet modifies the breadth "
        "and composition of that response without supporting generalized immune suppression as the principal "
        "explanation for diet-dependent mortality."
    )
    replace_text(p[45], abstract)

    replace_text(p[55], "2. Materials and methods", "Heading 1")
    replace_text(p[57], "2.1 Experimental design", "Heading 3")
    replace_text(
        p[58],
        "We used a 3 x 2 factorial design comprising three isocaloric diets and two inoculation treatments. "
        "The diets contained 33%, 50%, or 83% protein as a proportion of total dietary macronutrient, "
        "corresponding to whole-diet protein:carbohydrate percentages of 14:28, 21:21, and 35:7, respectively. "
        "Individual fifth-instar female Schistocerca gregaria were randomly assigned to a diet and received "
        "either Metarhizium robertsii conidia in oil or sterile oil alone. The RNA-seq data set analyzed here "
        "comprised 44 individuals: diet 33 included seven controls and nine infected locusts, diet 50 included "
        "six controls and nine infected locusts, and diet 83 included seven controls and six infected locusts "
        "(20 controls and 24 infected locusts in total). All 44 libraries were retained in the primary analysis."
    )

    replace_text(p[60], "2.2 Host, pathogen, and diets", "Heading 3")
    replace_text(p[61], "Study animals", "Heading 4")
    replace_text(
        p[62],
        "Fifth (final) instar female S. gregaria, three days post-molt, were obtained from a gregarious colony "
        "maintained at Arizona State University from laboratory lines originating at the University of "
        "Leicester (United Kingdom) and the University of Haifa-Oranim (Israel). Age was controlled during "
        "density rearing to reduce age-dependent transcriptional variation, and the three-day interval allowed "
        "cuticular hardening before inoculation. Colony conditions were 34 degrees C during the day, 25 degrees C "
        "at night, 20-50% relative humidity, and a 14 h:10 h light:dark cycle. Colony locusts received organic "
        "romaine lettuce, wheat grass, and wheat bran. Experimental locusts were housed individually in "
        "ventilated plastic containers (18.9 x 13.5 x 9.5 cm) with food and water provided ad libitum."
    )
    replace_text(p[63], "Fungal culture and inoculation", "Heading 4")
    replace_text(
        p[64],
        "Locusts were challenged with M. robertsii isolate DWR2009 (also designated ARSEF 10343), originally "
        "isolated from soil in the United States. Conidia were produced by biphasic solid-state fermentation "
        "following Jaronski and Jackson (2012), suspended in Orchex 792 paraffinic oil, and adjusted to 1 x "
        "10^9 conidia mL-1. Concentration was determined with a hemocytometer and viability was confirmed by "
        "germination on potato dextrose agar at 25-26 degrees C. Infected locusts received 1 microliter of the "
        "conidial suspension at the hind-leg femur-trochanter junction using a Hamilton Gastight syringe with "
        "a repeating dispenser and blunt-ended PTFE needle. Controls received 1 microliter of sterile Orchex "
        "792 oil and were handled identically."
    )
    replace_text(p[65], "Artificial diets", "Heading 4")
    replace_text(
        p[66],
        "Chemically defined diets followed Dadd (1961), Simpson and Abisgold (1985), and the formulation used "
        "by Zembrzuski et al. (2023). Each diet contained 42% macronutrient by dry mass and was supplemented "
        "with Wesson's salts (2.4%), cholesterol (0.5%), linoleic acid (0.5%), ascorbic acid (0.3%), and a "
        "multivitamin mixture (0.2%). Protein was supplied as a 3:1:1 mixture of casein, peptone, and albumen; "
        "carbohydrate was supplied as equal parts sucrose and dextrin. Diet and water were replenished as needed."
    )

    replace_text(p[68], "2.3 Tissue sampling and RNA extraction", "Heading 3")
    blank(p[69])
    replace_text(
        p[70],
        "Fat body was collected 96 h after inoculation because this tissue integrates insect immunity, "
        "nutrient storage, and intermediary metabolism. A single experimenter performed all dissections under "
        "sterile, RNase-free conditions. Instruments were flame sterilized and sequentially treated with 10% "
        "bleach, distilled water, 70% ethanol, and RNaseZap before each dissection. Locusts were surface "
        "sterilized with 70% ethanol, and abdominal fat body was removed while avoiding the gut, trachea, and "
        "heart. Tissue was transferred immediately to nuclease-free tubes, flash-frozen in liquid nitrogen, "
        "and stored at -80 degrees C."
    )
    replace_text(
        p[71],
        "Total RNA was extracted on a Promega Maxwell RSC instrument with the simplyRNA Tissue Kit according "
        "to the manufacturer's protocol. RNA concentration was measured with a Qubit 4 Fluorometer and the "
        "Qubit RNA HS Assay. RNA integrity and dissection quality were assessed before library preparation."
    )

    replace_text(p[72], "2.4 Library preparation and sequencing", "Heading 3")
    replace_text(
        p[73],
        "Stranded total-RNA libraries were prepared at Texas A&M University with the Illumina Stranded Total "
        "RNA kit and Ribo-Zero rRNA depletion. Pooled libraries were sequenced by Novogene on an Illumina "
        "NovaSeq X Plus 25B platform as 150-bp paired-end reads, targeting approximately 45 million read pairs "
        "per library."
    )
    replace_text(
        p[74],
        "Residual rRNA was quantified from assigned gene-level fragments before filtering. The median residual "
        "rRNA fraction was 7.9% per library (range 3.3-41.7%). All 12,402 gene identifiers in the project rRNA "
        "exclusion list were removed before expression filtering, normalization, ordination, and differential "
        "expression testing."
    )

    replace_text(p[75], "2.5 Read processing, competitive mapping, and gene counting", "Heading 3")
    replace_text(
        p[76],
        "RNA-seq preprocessing was executed on a high-performance computing cluster with a run-scoped "
        "Snakemake workflow. Paired reads were processed with fastp using automatic paired-end adapter "
        "detection, removal of the first two bases of each read, and a minimum retained length of 50 bp. JSON "
        "and HTML quality reports were retained for each library."
    )
    replace_text(
        p[77],
        "Trimmed reads were aligned in two-pass mode with STAR v2.7.10b. The primary branch used a competitive "
        "reference comprising the S. gregaria assembly GCF_023897955.1 and the M. robertsii ARSEF 23 assembly "
        "GCF_000187425.2. The latter is a close reference genome rather than the exact challenge isolate "
        "DWR2009/ARSEF 10343. A host-only alignment branch was retained as a mapping sensitivity analysis. "
        "Genome indices used sjdbOverhang = 149 and alignIntronMax = 2,500,000, selected after examination of "
        "Schistocerca intron lengths. Coordinate-sorted BAMs and transcriptome BAMs were retained and indexed "
        "with samtools v1.21."
    )
    replace_text(
        p[78],
        "Gene-level fragments were quantified with featureCounts (Subread) using paired-end fragment counting "
        "(--countReadPairs), feature types transcript and exon (-t transcript,exon), and gene_id as the grouping "
        "attribute (-g gene_id). This span-inclusive definition assigns fragments overlapping either annotated "
        "feature class to a gene and can capture signal consistent with incompletely processed or nascent RNA; "
        "an exon-only matrix was analyzed separately as a count-definition sensitivity check."
    )
    replace_text(
        p[79],
        "Sequence placement was assigned from the official assembly metadata rather than scaffold-name patterns. "
        "The primary manuscript analysis was restricted to non-rRNA genes on the main assembled chromosomes. "
        "All-scaffold results were retained as a placement sensitivity analysis, and competitive-host and "
        "host-only placed-chromosome DEG totals were compared to assess mapping dependence."
    )

    replace_text(p[80], "2.6 Differential expression and sensitivity analyses", "Heading 3")
    replace_text(
        p[81],
        "DESeq2 was applied to raw gene counts after rRNA and sequence-placement filtering. Genes entered a model "
        "when they had at least 10 counts in at least three libraries. Automatic count replacement was disabled "
        "(minReplicatesForReplace = Inf), and Cook's-distance filtering was disabled for planned contrasts "
        "(cooksCutoff = FALSE), so all 44 samples contributed to the primary estimates. A gene was called "
        "differentially expressed at a Benjamini-Hochberg adjusted P value below 0.05 and an absolute log2 fold "
        "change of at least 1. DESeq2 estimated sample-specific size factors from the integer count matrix to "
        "account for differences in sequencing depth and library composition during model fitting. "
        "Variance-stabilized values derived from the normalized counts were used only for PCA and heatmaps, "
        "not as input to differential-expression testing."
    )
    replace_text(
        p[82],
        "A six-level diet-by-treatment group model (~ 0 + group) generated direct within-diet infection contrasts "
        "and diet contrasts separately among controls and infected locusts. A second model (~ Diet + Treatment) "
        "estimated the average infection effect across diets, and a full factorial model (~ Diet + Treatment + "
        "Diet:Treatment) tested whether infection responses differed among diets. For diet contrasts, diet 50 "
        "served as the reference for comparisons with diets 33 and 83, and diet 33 served as the reference for "
        "the 83-versus-33 comparison. Positive log2 fold changes indicate higher expression in the named "
        "numerator condition."
    )
    replace_text(
        p[83],
        "Library size, variance-stabilized PCA, sample-distance heatmaps, and heatmaps of highly variable genes "
        "were examined for sample-level quality control. Centroids were calculated for each diet-by-treatment "
        "group, and lines joining control and infected centroids summarized the direction of the infection shift "
        "within each diet."
    )
    replace_text(
        p[84],
        "Sample 1044 was excluded before biological modelling because of its technical mapping failure. Samples "
        "1007, 1036, 1037, and 1039 passed sequencing and mapping quality control and were retained despite their "
        "positions in exploratory ordination; no PCA-defined sample was removed. Additional sensitivities compared exon-only with "
        "transcript-plus-exon counts and main-chromosome with all-scaffold analyses. None of these sensitivity "
        "branches replaced the all-sample, transcript-plus-exon, main-chromosome primary analysis."
    )

    replace_text(p[86], "2.7 Gene classification and functional enrichment", "Heading 3")
    replace_text(
        p[87],
        "Reference annotation and local crosswalk tables classified genes as protein coding, long non-coding RNA "
        "(lncRNA), or other/unknown. Because every listed rRNA locus was removed before modeling, the latter "
        "category contains non-rRNA features or genes without a resolved biotype."
    )
    replace_text(
        p[88],
        "Functional descriptions, Gene Ontology (GO) terms, and KEGG orthology/pathway assignments were obtained "
        "from the project eggNOG annotation and linked to gene identifiers through CDS records in the reference "
        "GFF. KEGG results were shown in a eukaryote-inclusive view and were interpreted as orthology-based, "
        "hypothesis-generating annotations rather than direct evidence that an insect pathway was activated."
    )
    replace_text(
        p[89],
        "Over-representation tests were performed with clusterProfiler::enricher using the 14,132 expressed, "
        "non-rRNA main-chromosome genes entering the primary DESeq2 analysis as the background. P values were "
        "adjusted by the Benjamini-Hochberg method. For the two expression clusters, enrichment foregrounds were "
        "restricted to genes differentially expressed in at least one within-diet infected-versus-control "
        "contrast. GO Biological Process, Cellular Component, and Molecular Function terms and eukaryote-compatible "
        "KEGG pathways were summarized separately. Term-to-gene reverse-index tables were retained so each "
        "enrichment result could be traced to its contributing genes."
    )

    replace_text(p[91], "2.8 Curated physiological candidate genes", "Heading 3")
    replace_text(
        p[92],
        "A curated candidate list was used to examine immune recognition and signaling, humoral immunity, cellular "
        "immunity, detoxification, protein synthesis and translation, nutrient sensing and TOR/insulin signaling, "
        "nitrogen and amino-acid metabolism, and storage hexamerins. The primary placed-chromosome analysis "
        "retained 99 candidates. For each contrast, candidates were classified as upregulated, downregulated, or "
        "not significant using the same DEG thresholds as the genome-wide analysis. Chi-squared tests compared "
        "these proportions among diets within each functional category. The all-scaffold candidate set was "
        "retained as a placement sensitivity analysis."
    )

    replace_text(
        p[90],
        "Genes significant in any primary contrast were clustered from row-standardized variance-stabilized "
        "expression. The gene dendrogram supported two broad expression programs. Genes were ordered within each "
        "cluster without averaging individual-gene expression, preserving sample-level heterogeneity in the heatmap."
    )

    # Add the independent phase-reference comparison and fungal-RNA proxy before Results.
    phase_heading = p[94].insert_paragraph_before(
        "2.9 Assessment of phase change as a potential confounder", style="Heading 3"
    )
    phase_text = p[94].insert_paragraph_before(
        "Because the infection protocol required locusts to be maintained individually under semi-isolated conditions "
        "for 96 h, we assessed whether an early density-dependent phase response could contribute to heterogeneous "
        "fat-body expression and therefore act as a potential confounding factor. The union of genes significant "
        "in at least one within-diet infection contrast was compared with independent solitarious-versus-gregarious "
        "DEG sets from S. gregaria head and thorax. The reference study also used female final-instar nymphs sampled "
        "three days post-molt and the same total-RNA, rRNA-depletion library-preparation workflow. Head and thorax "
        "were evaluated separately; genes detected in both tissues were treated as the most conservative cross-tissue "
        "phase-associated signature, while one-tissue overlaps were retained as tissue-dependent phase evidence. "
        "These overlaps identify genes whose expression could reflect infection, an early response to reduced social "
        "density, or both; they do not establish which process caused the response. Genes absent from both sets were "
        "considered unsupported by the available phase references, rather than proven infection-specific. This "
        "comparison assessed phase change as a possible confounder but did not establish that a phase transition "
        "occurred in the experimental fat body."
    )
    phase_heading.paragraph_format.keep_with_next = True
    phase_text.paragraph_format.space_after = Pt(8)

    fungal_heading = p[94].insert_paragraph_before("2.10 Fungal alignment proxy", style="Heading 3")
    fungal_text = p[94].insert_paragraph_before(
        "Uniquely aligned fragments assigned to the fungal component of the competitive reference were summarized "
        "as counts and as a percentage of uniquely aligned competitive-reference fragments. Because the mapping "
        "reference was not strain matched and fungal load was not confirmed post mortem by qPCR or culture, this "
        "quantity was treated only as RNA-seq evidence of fungal transcriptional signal. A parallel qPCR validation "
        "would have been preferable, but no matching tissue remained for that assay."
    )
    fungal_heading.paragraph_format.keep_with_next = True
    fungal_text.paragraph_format.space_after = Pt(8)

    # Results are rewritten around the current placed-chromosome outputs.
    replace_text(p[95], "3.1 Global fat body expression and fungal RNA signal", "Heading 3")
    replace_text(
        p[96],
        "The first two principal components of variance-stabilized expression explained 27.7% and 12.4% of total "
        "variance, respectively (Fig. 1). Control samples occupied the negative portion of PC1, whereas most "
        "infected samples shifted toward positive PC1 values. The control-to-infected centroid moved in the same "
        "general direction for each diet, although the hulls overlapped and within-group dispersion differed, "
        "particularly among infected locusts. Infection therefore represented the dominant common axis of "
        "variation, with diet and inter-individual heterogeneity modifying the position and breadth of each group."
    )
    replace_text(
        p[97],
        "Fungal alignment signal was low in most fat-body libraries but highly heterogeneous among infected "
        "individuals. Median fungal percentages were 0.00138%, 0.00097%, and 0.00084% in infected diets 33, 50, "
        "and 83, respectively, while a small number of infected libraries reached 0.64-3.13%. Controls remained "
        "near background. Because these values came from competitive RNA-seq alignment rather than a strain-matched "
        "load assay, they were used to describe detectable fungal RNA, not to classify infection success."
    )
    add_picture(p[98], FIGURES["figure1"], 6.2)
    add_caption(
        p[99],
        "Figure 1.",
        "Fat body transcriptomes shift with fungal infection across the three diets. PCA used variance-stabilized "
        "counts from expressed, non-rRNA genes on the main S. gregaria chromosomes. Points represent individual "
        "libraries, shapes indicate treatment, colors indicate diet, shaded polygons show the observed breadth of "
        "each diet-by-treatment group, open circles mark centroids, and lines connect the control and infected "
        "centroids within each diet."
    )

    replace_text(p[100], "3.2 Differential expression across infection and diet contrasts", "Heading 3")
    replace_text(
        p[101],
        "The average infected-versus-control contrast identified 1,691 DEGs across diets, including 893 genes "
        "higher in infected locusts and 798 higher in controls (Fig. 2). Within diets, infection altered 1,016 "
        "genes in diet 33 (552 higher in infected and 464 higher in controls), 2,009 genes in diet 50 (1,263 and "
        "746), and 1,030 genes in diet 83 (658 and 372). Protein-coding genes formed the majority of each response, "
        "but lncRNAs contributed 51, 132, and 62 DEGs in diets 33, 50, and 83, respectively."
    )
    replace_text(
        p[102],
        "Direct diet effects were much smaller among controls: pairwise contrasts yielded 15 DEGs between diets "
        "33 and 50, 60 between diets 83 and 33, and 59 between diets 83 and 50. Diet contrasts among infected "
        "locusts yielded 303, 86, and 59 DEGs for the same comparisons. The particularly large infected 33-versus-50 "
        "contrast was strongly directional, with 290 genes higher in diet 50 and 13 higher in diet 33. Across the "
        "three within-diet infection responses, 467 genes were shared by all diets, whereas 211, 1,069, and 320 "
        "were unique to diets 33, 50, and 83, respectively."
    )
    add_picture(p[104], FIGURES["figure2"], 6.15)
    add_caption(
        p[105],
        "Figure 2.",
        "Differential-expression burden and overlap in the all-sample, main-chromosome analysis. The upper panels "
        "show coding, lncRNA, and other/unknown DEGs by direction for the average infection response, infection "
        "within each diet, diet contrasts among controls, and diet contrasts among infected locusts. Red values are "
        "higher in the named numerator condition and blue values are higher in the denominator. The lower-left Venn "
        "diagram partitions the three within-diet infection DEG sets. The lower-right diagram compares their union "
        "with external head and thorax solitarious-versus-gregarious DEG sets. Genes shared with both tissues provide "
        "the most conservative cross-tissue phase-associated subset; one-tissue overlaps provide tissue-dependent "
        "phase context. Absence from both sets indicates a lack of support in these references, not proof of infection "
        "specificity. Annotated rRNA genes were absent from every set."
    )
    blank(p[103])
    blank(p[106])
    blank(p[107])

    replace_text(p[110], "3.3 Infection-associated expression programs and functional enrichment", "Heading 3")
    replace_text(
        p[111],
        "The union of genes significant in the primary contrast set comprised 2,949 main-chromosome genes. "
        "Hierarchical clustering resolved two broad programs without averaging gene-level expression (Fig. 3). "
        "Cluster 1 contained 1,086 genes and generally showed higher expression in controls and lower expression "
        "after infection. Cluster 2 contained 1,863 genes and showed the reciprocal pattern. Most genes in both "
        "clusters were associated with an infection-within-diet contrast; smaller subsets reflected diet effects "
        "or membership in more than one contrast context."
    )
    replace_text(
        p[112],
        "Among within-diet infection DEGs, cluster 1 was enriched for chitin-based cuticle development, mitotic "
        "cell-cycle and spindle functions, oxidoreductase activity, broad metabolism, cytochrome P450-associated "
        "xenobiotic metabolism, and insect hormone biosynthesis. Cluster 2 was enriched for response to bacteria "
        "and stimulus, extracellular and plasma-membrane compartments, peptidoglycan metabolism and binding, "
        "transmembrane transport, and Toll and Imd signaling. These reciprocal programs are consistent with "
        "reduced investment in cuticle, cell-cycle, and selected metabolic processes alongside increased expression "
        "of extracellular defense and recognition functions during infection."
    )
    replace_text(
        p[113],
        "Only 325 cluster-1 and 396 cluster-2 foreground genes had GO annotations, and 264 and 318, respectively, "
        "had KEGG assignments. The enrichment results therefore summarize the annotated subset rather than every "
        "DEG. Fungal alignment percentages displayed above the heatmap emphasized that the host expression "
        "programs occurred across a wide range of fungal RNA signals."
    )
    add_picture(p[108], FIGURES["figure3"], 5.15)
    add_caption(
        p[109],
        "Figure 3.",
        "Main-chromosome DEG expression programs and functional enrichment. (A) Gene-level row-z-scored "
        "variance-stabilized expression for 2,949 DEGs, separated into two dendrogram-supported clusters and "
        "ordered with controls before infected samples. The upper bar shows the percentage of unique competitive "
        "alignments assigned to M. robertsii; values above the 0.1% display range are labeled. (B) GO terms enriched "
        "among within-diet infection DEGs in each cluster. (C) Eukaryote-inclusive KEGG pathways enriched in the "
        "same foregrounds. Point size indicates the number of cluster genes in a term. Cluster-1 enrichment values "
        "are mirrored to the left only to separate the panels; the sign does not represent regulation direction."
    )

    replace_text(p[114], "3.4 Curated immune and physiological candidate genes", "Heading 3")
    replace_text(
        p[115],
        "The main-chromosome curated analysis retained 99 genes across eight immune and protein-anabolism "
        "categories (Fig. 4). Within each diet, infection increased a substantial fraction of genes involved in "
        "innate immune recognition and signaling and humoral immunity. Responses in cellular immunity and "
        "detoxification were more variable, and several genes in these small categories were not significant."
    )
    replace_text(
        p[116],
        "Protein-anabolism categories contained both increased and decreased responses. Protein synthesis and "
        "translation included a small upregulated fraction and a larger downregulated fraction in every diet. "
        "Nutrient-sensing and nitrogen/amino-acid categories tended to contain more upregulated genes in diet 50 "
        "than in diets 33 or 83, whereas storage hexamerins were mostly downregulated in diet 50 and not significant "
        "in diets 33 and 83."
    )
    replace_text(
        p[117],
        "Despite these descriptive tendencies, chi-squared tests detected no significant among-diet difference in "
        "the proportions of upregulated, downregulated, and non-significant genes for any functional category. "
        "Across-diet contrasts among infected locusts likewise contained relatively few significant curated genes. "
        "The curated analysis therefore identifies candidate pathways for follow-up but does not support a "
        "category-wide diet effect."
    )
    replace_text(
        p[118],
        "Restricting the curated list to main chromosomes reduced it from 163 all-scaffold candidates to 99. This "
        "placement filter was applied consistently to the primary PCA, differential-expression, clustering, "
        "enrichment, and curated analyses; the all-scaffold results are retained as a supplementary sensitivity."
    )
    blank(p[119])
    add_picture(p[120], FIGURES["figure4"], 6.15)
    blank(p[121])
    add_caption(
        p[122],
        "Figure 4.",
        "Curated immune and protein-anabolism candidate genes on the main S. gregaria chromosomes. Stacked bars "
        "show the percentage of genes that were upregulated, downregulated, or not significant. (A, C) Infection "
        "responses within each diet. (B, D) Pairwise diet contrasts among infected locusts, oriented as higher- "
        "versus lower-protein diets. The number below each functional heading is the number of retained candidate "
        "genes. N.S.D. indicates that the upregulated/downregulated/non-significant proportions did not differ "
        "significantly among the displayed comparisons."
    )
    blank(p[123])

    # Discussion: retain the manuscript's voice while removing claims not supported by the current tests.
    replace_text(
        p[125],
        "Metarhizium infection was the principal driver of transcriptional variation in the desert locust fat "
        "body, whereas dietary macronutrient balance modified the breadth and composition of this response. The "
        "control-to-infected centroid moved in a common direction across diets, and the average infection contrast "
        "identified 1,691 DEGs. However, the within-diet response was almost twice as large in the balanced diet "
        "(2,009 DEGs) as in the carbohydrate-biased (1,016) or protein-biased (1,030) diets. Direct diet contrasts "
        "were comparatively small in controls and larger, but still more limited, among infected locusts. These "
        "patterns support a dominant infection response whose expression is conditioned by nutritional state."
    )
    replace_text(
        p[126],
        "The 467 genes shared among all three within-diet infection responses represent a conserved fat body core, "
        "whereas the 1,069 genes unique to diet 50 account for much of the larger response under the balanced diet. "
        "The two expression programs provide a useful functional summary: one program decreased during infection "
        "and was associated with cuticle, cell-cycle, and selected metabolic processes, while the other increased "
        "and was associated with extracellular defense, peptidoglycan recognition, transport, and Toll/Imd "
        "signaling. Infection therefore involved both immune activation and broad reallocation away from structural "
        "and proliferative processes."
    )
    replace_text(
        p[127],
        "The curated analysis refined this interpretation but also imposed an important constraint. Immune genes "
        "were commonly upregulated within every diet, and protein-anabolism genes showed mixed responses, yet none "
        "of the functional categories differed significantly among diets in their proportions of upregulated, "
        "downregulated, and non-significant genes. The transcriptome therefore does not support generalized immune "
        "suppression as the main explanation for the shorter survival previously observed on balanced and "
        "protein-biased diets (Tahir et al., 2026). Diet may instead alter particular genes, response consistency, "
        "or physiological resource availability without producing a category-wide shift detectable in this data set."
    )
    replace_text(
        p[128],
        "Earlier physiological and metabolomic work in this system showed that higher-protein diets increased "
        "hemolymph protein and branched-chain amino-acid availability and were associated with accelerated mortality "
        "and fungal development, despite elevated phenoloxidase activity (Tahir et al., 2026; Tahir et al., "
        "submitted). The present mixed responses in nutrient sensing, amino-acid metabolism, translation, and "
        "storage genes are compatible with infection-driven metabolic reorganization, but the non-significant "
        "category tests do not establish that diet 50 or diet 83 produced a uniformly more anabolic fat body. "
        "Combining these data is therefore more consistent with a resource-availability or tolerance hypothesis "
        "than with simple immune failure, while leaving the causal mechanism unresolved."
    )
    replace_text(
        p[129],
        "Fungal RNA provided an additional, but limited, view of disease progression. Most infected fat-body "
        "libraries contained little uniquely aligned fungal signal, whereas several individuals contained markedly "
        "higher values. This heterogeneity could reflect differences in infection progression, fungal localization, "
        "or sampling of a tissue in which fungal biomass was low at 96 h. The competitive reference used ARSEF 23 "
        "rather than the challenge isolate DWR2009/ARSEF 10343, and no matching tissue remained for qPCR or culture. "
        "Consequently, these alignments should not be interpreted as a validated fungal-load measurement or used to "
        "reclassify infected individuals."
    )
    replace_text(
        p[130],
        "The marked expression of some DEG groups in only a few individuals raised the possibility that heterogeneous "
        "responses to 96 h of semi-isolation contributed to the apparent infection signal. Genes shared by the "
        "infection response and both reference "
        "tissues provide the strongest cross-tissue phase-associated subset, whereas overlap with only head or thorax "
        "may reflect tissue-dependent phase associations or shared stress responses. Infection DEGs absent from both "
        "sets lack support in the available phase references, but cannot be regarded as definitively infection-specific "
        "because the independent study sampled different tissues. Thus, phase-related expression may confound a "
        "subset of infection-responsive genes, but the comparison neither demonstrates a completed phase transition "
        "nor shows that solitarization explains the overall fat-body infection response."
    )
    replace_text(
        p[131],
        "Sequence placement was another important source of uncertainty. A large historical DEG signal occurred on "
        "unplaced scaffolds, especially in the infected diet-33-versus-diet-50 contrast. Taxonomic, cross-genome, "
        "and homology audits did not justify treating every unplaced locus as contamination, but their provenance "
        "was less secure than that of genes on assembled chromosomes. Restricting the primary analysis to main "
        "chromosomes therefore provides a conservative host interpretation, while preserving all-scaffold results "
        "for sensitivity and future investigation rather than deleting those loci."
    )
    replace_text(
        p[132],
        "Functional enrichment in a non-model insect remains constrained by annotation coverage and orthology. Only "
        "a subset of genes in either expression cluster carried GO or KEGG assignments, and broad KEGG pathway names "
        "may reflect conserved orthologous functions rather than a locust-specific pathway definition. The use of "
        "transcript-plus-exon counts also increased sensitivity relative to exon-only counting, particularly for "
        "within-diet infection contrasts, but it can include signal consistent with nascent or incompletely processed "
        "RNA. The count-definition, sequence-placement, and sample-retention analyses are therefore reported as "
        "explicit sensitivities. Functional enrichment should be considered hypothesis-generating until supported by "
        "protein, metabolite, or functional assays."
    )
    replace_text(
        p[133],
        "Together, these results support a model in which M. robertsii induces a conserved and extensive fat body "
        "response, while dietary macronutrient balance changes the scale and gene composition of that response. The "
        "balanced diet produced the broadest transcriptional response, but the curated immune and protein-anabolism "
        "tests did not identify a significant category-wide diet effect. Diet-dependent mortality is therefore "
        "unlikely to be explained by generalized immune suppression alone. Future experiments combining matched "
        "fungal-load qPCR, proteomics, metabolite flux, and pathway-specific perturbation will be required to resolve "
        "whether host resource allocation, pathogen nutrition, or tolerance drives the observed survival differences."
    )

    # Replace old supplementary figure placeholders with a focused sensitivity package.
    replace_text(
        p[218],
        "This supplementary information reports sample-retention, sequence-placement, count-definition, fungal-RNA, "
        "and curated-target sensitivities supporting the primary all-sample, main-chromosome analysis."
    )
    cover_heading = p[236]
    for old_paragraph in list(doc.paragraphs[219:236]):
        remove_paragraph(old_paragraph)

    supplement_captions = [
        (
            "figureS1", 6.1, "Figure S1.",
            "Mapping-quality basis for excluding sample 1044. All 45 sequenced libraries are shown for unique mapping "
            "and unique plus accepted multimapping. Sample 1044 failed technical host-alignment QC and was excluded; "
            "all other libraries were retained in the biological analyses."
        ),
        (
            "figureS2", 6.15, "Figure S2.",
            "Chromosomal localization of differential expression. Main-chromosome DEGs are shown by genomic "
            "position and contrast context, providing a genome-wide view of whether signals are broadly distributed "
            "or concentrated in particular chromosome regions."
        ),
        (
            "figureS3", 6.15, "Figure S3.",
            "Count-definition sensitivity. Numbers of significant genes obtained with exon-only counts are compared "
            "with counts assigned from transcript and exon features and grouped by gene_id. Both matrices were "
            "depleted of annotated rRNA and restricted to the same primary chromosome scope."
        ),
        (
            "figureS4", 6.15, "Figure S4.",
            "Competitive-alignment evidence for fungal RNA. Full-range and zoomed views show uniquely aligned "
            "M. robertsii fragments and their percentage of unique competitive-reference alignments. Horizontal "
            "control means provide a background reference. These values are RNA-seq proxies, not qPCR-validated "
            "fungal loads."
        ),
        (
            "figureS5", 6.15, "Figure S5.",
            "Sequence-placement sensitivity for the curated candidate analysis. Results from all annotated scaffolds "
            "are compared with the primary main-chromosome subset, showing which changes arise from excluding "
            "unplaced or otherwise non-primary assembly sequences."
        ),
    ]

    for key, width, label, caption in supplement_captions:
        clean_name = f"{key}.png"
        copied = SUPPLEMENT_DIR / clean_name
        shutil.copy2(FIGURES[key], copied)
        add_supplement_before(cover_heading, copied, width, label, caption)

    # Bring the cover-letter summary into agreement with the revised manuscript.
    replace_text(
        p[245],
        "Our study complements earlier physiological measurements by testing the interactive effects of diet and "
        "Metarhizium infection on fat body transcription. Infection produced a large conserved response, with a "
        "467-gene core shared among all diets, while the balanced diet produced the broadest within-diet response. "
        "Curated immune and protein-anabolism categories showed biologically informative tendencies but no "
        "significant category-wide differences among diets. Together with the earlier hemolymph and survival data, "
        "these results argue against generalized immune suppression as the sole mechanism and instead motivate "
        "tests of host resource allocation, pathogen nutrition, and disease tolerance."
    )

    # Keep figures and captions together where possible and use journal-friendly body spacing.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.lower() in {"normal", "body text"}:
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.08

    doc.core_properties.title = (
        "Fat body transcriptional responses to dietary macronutrient balance and Metarhizium infection"
    )
    doc.core_properties.comments = (
        "Clean CBP Part D revision generated from current all-sample, transcript-plus-exon, rRNA-depleted, "
        "main-chromosome analysis outputs."
    )
    doc.save(OUTPUT_DOCX)

    # A separate review copy highlights revised passages without marking the clean manuscript.
    review = Document(OUTPUT_DOCX)
    review_paragraphs = review.paragraphs
    note = review_paragraphs[0].insert_paragraph_before(
        "REVIEW COPY: passages revised from the supplied manuscript are highlighted in yellow."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].bold = True
    note.runs[0].font.size = Pt(10)

    highlight = False
    highlight_single_starts = (
        "Fungal infection strongly reshaped",
        "Diet 50 produced the broadest",
        "A core of 467 infection-responsive",
        "Curated immune and anabolic",
        "Dietary macronutrient balance can alter disease outcome",
        "This supplementary information reports sample-retention",
        "Our study complements earlier physiological measurements",
    )
    for paragraph in review.paragraphs:
        text = paragraph.text.strip()
        if text in {"2. Materials and methods", "3. Results", "4. Discussion"}:
            highlight = True
        if text in {"Acknowledgements", "Draft cover letter"}:
            highlight = False
        should_highlight = highlight or text.startswith(highlight_single_starts)
        if should_highlight:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Highlight the replacement supplementary figure package up to the cover letter.
    in_supplements = False
    for paragraph in review.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("This supplementary information reports sample-retention"):
            in_supplements = True
        if text == "Draft cover letter":
            in_supplements = False
        if in_supplements:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    review.core_properties.comments = (
        "Review copy of the CBP Part D revision. Yellow highlighting identifies rewritten passages."
    )
    review.save(OUTPUT_REVIEW_DOCX)
    print(OUTPUT_DOCX)
    print(OUTPUT_REVIEW_DOCX)


if __name__ == "__main__":
    main()
